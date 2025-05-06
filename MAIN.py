#!/usr/bin/env python3
"""
AGT Price Tag Transformer
--------------------------
This script reads and processes an Excel file, applies various filters,
generates dynamic DOCX files with product labels/tags using multiple DOCX libraries,
and provides a Tkinter GUI for user interaction.
"""
import sys, os, platform, subprocess, re, datetime, math, traceback, logging
import concurrent.futures
from io import BytesIO
from copy import deepcopy
from xml.sax.saxutils import unescape
import tkinter as tkmod
from tkinter import filedialog, messagebox, ttk
import pandas as pd
import numpy as np
from functools import lru_cache
from pathlib import Path
import io
import threading, urllib.request, json
import urllib.request
import json
from difflib import SequenceMatcher




# ------------------ Third-Party DOCX Imports ------------------
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docxcompose.composer import Composer
import docxcompose
import os, datetime
from tkinter import simpledialog
import urllib.request, json


TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__),
    "templates",
    "inventory_slip.docx"
)

INVENTORY_SLIP_TEMPLATE = os.path.join(
    os.path.dirname(__file__),
    "templates",
    "inventory_slip.docx"
)



LOG_PATH = os.path.join(
    os.path.expanduser("~"),
    "Downloads",
    "lineage_change_log.csv"
)


import os, docxcompose
docxcompose_templates = (os.path.join(os.path.dirname(docxcompose.__file__), "templates"), "docxcompose/templates")

logging.basicConfig(level=logging.DEBUG, format="%(asctime)s %(levelname)s: %(message)s")
logging.debug("Application starting...")

from concurrent.futures import ThreadPoolExecutor
_IO_POOL = ThreadPoolExecutor(max_workers=1)          # single worker for I/O

from concurrent.futures import ThreadPoolExecutor
executor = ThreadPoolExecutor(max_workers=4)

def _add_cat_value(series, value):
    """Add *value* to a Categorical column if it isn't present."""
    if hasattr(series.dtype, "categories") and value not in series.cat.categories:
        return series.cat.add_categories([value])
    return series

_UPDATING_FILTERS = False

# Global variable to track which canvas is active.
# Global variables
current_canvas = None
available_canvas = None
selected_canvas  = None
available_tags_container = None
selected_tags_container = None
product_state_vars = {}
undo_stack = []  # For undo functionality
placeholder_img = None
print_vendor_back_var = None
SCALE_FACTOR = 1.0
WORD_WEIGHT = 5
json_matched_names = None



# near the top of your file, define the desired lineage order:
LINEAGE_COLOR_MAP = {
    "SATIVA":        "#E74C3C",
    "INDICA":        "#8E44AD",
    "HYBRID":        "#27AE60",
    "HYBRID/SATIVA": "#E74C3C",
    "HYBRID/INDICA": "#8E44AD",
    "CBD":           "#F1C40F",
    "MIXED":         "#2C3E50",
    "PARAPHERNALIA": "#FF69B4",
}

import os, platform

def save_docx_and_replace(doc, path):
    """
    Save a python‑docx Document to `path`. If Word has that exact file open,
    it will close it for you (no changes saved), then overwrite it.
    """
    try:
        doc.save(path)
    except PermissionError:
        system = platform.system()
        if system == "Windows":
            from win32com.client import Dispatch
            word = Dispatch("Word.Application")
            for d in word.Documents:
                if os.path.abspath(d.FullName).lower() == os.path.abspath(path).lower():
                    d.Close(False)
                    break
        elif system == "Darwin":
            name = os.path.basename(path)
            applescript = f'''
            tell application "Microsoft Word"
              close (every document whose name is "{name}") saving no
            end tell
            '''
            os.system(f"osascript -e '{applescript}'")
        else:
            raise
        doc.save(path)

def on_load_json_url(url: str):
    if not url.lower().startswith("http"):
        messagebox.showerror("Invalid URL", "Please paste a valid JSON URL.")
        return
    # fetch in a thread so the UI doesn’t freeze
    threading.Thread(target=_fetch_and_match, args=(url,), daemon=True).start()

from difflib import SequenceMatcher
import re

# compile once
_DIGIT_UNIT_RE = re.compile(r"\b\d+(?:g|mg)\b")
_NON_WORD_RE    = re.compile(r"[^\w\s-]")
_SPLIT_RE       = re.compile(r"[-\s]+")
# type‐override lookup
TYPE_OVERRIDES = {
    "all-in-one":      "vape cartridge",
    "rosin":           "concentrate",
    "mini buds":       "flower",
    "bud":             "flower",
    "pre-roll":        "pre-roll",
}

# on module load, build normalized‐desc/tokens cache
_sheet_cache = None
def _build_sheet_cache():
    global _sheet_cache
    df = global_df[
        global_df["Description"].notna() &
        ~global_df["Description"].str.lower().str.contains("sample", na=False)
    ]
    cache = []
    for idx, row in df.iterrows():
        desc = row["Description"]
        norm = _SPLIT_RE.sub(" ",
               _NON_WORD_RE.sub(" ",
               _DIGIT_UNIT_RE.sub("", desc.lower())
        )).strip()
        toks = set(norm.split())
        cache.append({
            "idx": idx,
            "brand": row["Product Brand"].lower(),
            "vendor": row["Vendor"].lower(),
            "ptype": row["Product Type*"].lower(),
            "norm": norm,
            "toks": toks,
        })
    _sheet_cache = cache

def _fetch_and_match(url: str):
    splash = show_splash2(root)
    global json_matched_names, _sheet_cache
    if _sheet_cache is None:
        _build_sheet_cache()

    try:
        # fetch JSON
        with urllib.request.urlopen(url) as resp:
            payload = json.loads(resp.read().decode())
        items = payload.get("inventory_transfer_items", [])

        # gather JSON brands/vendor
        json_brands = { itm.get("product_brand","").lower() for itm in items if itm.get("product_brand") }
        json_vendor = payload.get("from_license_name","").lower()

        # prefilter cache by brand/vendor
        pre = [
            r for r in _sheet_cache
            if (r["brand"] in json_brands) or (r["vendor"] == json_vendor)
        ]

        matched_idxs = set()

        # normalize + tokens helper
        def normalize(s: str):
            s = (s or "").lower()
            s = _DIGIT_UNIT_RE.sub("", s)
            s = _NON_WORD_RE.sub(" ", s)
            return _SPLIT_RE.sub(" ", s).strip()

        # for each JSON name
        for itm in items:
            raw = itm.get("product_name") or ""
            name_norm = normalize(raw)
            if not name_norm:
                continue
            name_toks = set(name_norm.split())

            # type override
            override = next(
                (ptype for kw, ptype in TYPE_OVERRIDES.items() if kw in name_norm),
                None
            )

            # work on a slice of `pre`
            bucket = [r for r in pre if (override is None or r["ptype"] == override)]

            # 1) substring
            for r in bucket:
                if r["norm"] in name_norm or name_norm in r["norm"]:
                    matched_idxs.add(r["idx"])

            # 2) token‐overlap ≥2
            for r in bucket:
                if len(name_toks & r["toks"]) >= 2:
                    matched_idxs.add(r["idx"])

            # 3) Jaccard ≥0.3
            for r in bucket:
                u = name_toks | r["toks"]
                if u and len(name_toks & r["toks"]) / len(u) >= 0.3:
                    matched_idxs.add(r["idx"])

            # 4) SequenceMatcher fallback on the normalized whole
            short, long = (name_norm, r["norm"]) if len(name_norm) < len(r["norm"]) else (r["norm"], name_norm)
            win = len(short)
            for i in range(len(long) - win + 1):
                if SequenceMatcher(None, long[i:i+win], short).ratio() >= 0.6:
                    matched_idxs.update(r["idx"] for r in bucket)
                    break

        # always include all prefiltered rows
        matched_idxs.update(r["idx"] for r in pre)

        final = sorted(global_df.loc[list(matched_idxs), "Product Name*"].tolist())
        json_matched_names = final

        if final:
            root.after(0, lambda: populate_available_tags(final))
        else:
            root.after(0, lambda: messagebox.showinfo("JSON Match", "No items matched."))

    except Exception:
        logging.exception("[_fetch_and_match] failed")
        err = traceback.format_exc()
        root.after(0, lambda: messagebox.showerror("Error", f"Failed to fetch/match JSON:\n{err}"))
    finally:
        # always destroy splash2 when done
        root.after(0, splash.destroy)



posabit_instructions = (
    "How to Obtain and Download Your Excel File from POSaBit\n\n"
    "1. Navigate to the POSaBit Inventory Page\n"
    "   Open your preferred web browser and log into your POSaBit account. Once logged in, navigate to the POSaBit → Inventory → Lots section. "
    "This is where you will be able to view all available inventory lots.\n\n"
    
    "2. Set Up Your Filters\n"
    "   On the left-hand side of the screen, you will see a filter sidebar. You need to apply the following filters to display only the relevant lots:\n"
    "       • Status: Change the status filter to \"Active\" so that only active inventory items are shown.\n"
    "       • Quantity On Hand: Adjust the filter to show only items with a Quantity On Hand above 0. "
    "This ensures you are only downloading items that are currently in stock.\n\n"
    
    "3. Run the Search\n"
    "   Once you have set the filters, click the \"Search\" button. This action will refresh the list of lots to display only those that match your filter criteria (active items with available quantity).\n\n"
    
    "4. Download Your Excel File\n"
    "   After your search results have been updated, locate the blue Download arrow button and click it. "
    "This will download an Excel file containing your filtered data.\n\n"
    
    "5. Upload the Excel File and Select a Template\n"
    "   Return to this application and use the Upload button provided to select and upload the Excel file you just downloaded. "
    "Once uploaded, choose the appropriate template for generating product labels or inventory slips."
)


# ------------------ Helper Functions ------------------
def set_current_canvas(event, canvas):
    global current_canvas
    current_canvas = canvas

def clear_current_canvas(event):
    global current_canvas
    current_canvas = None

def wrap_with_marker(text, marker):
    safe_text = str(text).replace('&', '&amp;')
    return f"{marker.upper()}_START{safe_text}{marker.upper()}_END"

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS  # set by PyInstaller
    except Exception:
        base_path = os.path.abspath(".")
    full_path = os.path.join(base_path, relative_path)
    print(f"Loading resource from: {full_path}")
    return full_path

def open_file(file_path):
    if not os.path.exists(file_path):
        logging.error("File not found: %s", file_path)
        return
    try:
        if platform.system() == "Windows":
            os.startfile(file_path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["/usr/bin/open", "-a", "Microsoft Word", file_path])
        else:
            subprocess.Popen(["xdg-open", file_path])
    except Exception as e:
        logging.error("Error opening file: %s", e)


def get_default_upload_file():
    """
    Looks for files in the Downloads folder that start with "A Greener Today" and end with ".xlsx".
    Returns the full path of the most recently modified file, or None if no matching file is found.
    """
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    files = []
    
    # Scan the Downloads directory for matching files.
    for f in os.listdir(downloads_dir):
        if f.startswith("A Greener Today") and f.lower().endswith(".xlsx"):
            full_path = os.path.join(downloads_dir, f)
            files.append(full_path)
    
    if files:
        # Use the most recently modified file.
        latest_file = max(files, key=os.path.getmtime)
        return latest_file
    else:
        return None

# ------------------ Global Variables and Font Schemes ------------------
global_df = None  # DataFrame from Excel file
product_check_vars = {}  # (Legacy: not used for filtering labels anymore)
selected_tags_vars = {}   # Dictionary for items moved to Selected Tag List (key: product name)
available_tags_vars = {}    # Dictionary for available tag list
move_history = []


FONT_SCHEME_HORIZONTAL = {
    "DESC": {"base_size": 28, "min_size": 12, "max_length": 100},
    "PRIC": {"base_size": 38, "min_size": 20, "max_length": 20},
    "LINEAGE": {"base_size": 20, "min_size": 12, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 18, "min_size": 12, "max_length": 30},
    "THC_CBD": {"base_size": 12, "min_size": 10, "max_length": 50},
    "RATIO": {"base_size": 10, "min_size": 8, "max_length": 30},
    "WEIGHT": {"base_size": 18, "min_size": 10, "max_length": 20},
    "UNITS": {"base_size": 18, "min_size": 10, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 20, "min_size": 12, "max_length": 40}
}

FONT_SCHEME_VERTICAL = {
    "DESC": {"base_size": 23, "min_size": 12, "max_length": 100},
    "PRIC": {"base_size": 36, "min_size": 20, "max_length": 20},
    "LINEAGE": {"base_size": 18, "min_size": 12, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 18, "min_size": 12, "max_length": 30},
    "THC_CBD": {"base_size": 12, "min_size": 10, "max_length": 50},
    "RATIO": {"base_size": 8, "min_size": 10, "max_length": 30},
    "WEIGHT": {"base_size": 16, "min_size": 10, "max_length": 20},
    "UNITS": {"base_size": 16, "min_size": 10, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 20, "min_size": 12, "max_length": 40}
}

FONT_SCHEME_MINI = {
    "DESC": {"base_size": 20, "min_size": 8, "max_length": 100},
    "PRIC": {"base_size": 22, "min_size": 10, "max_length": 20},
    "LINEAGE": {"base_size": 10, "min_size": 8, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 10, "min_size": 8, "max_length": 30},
    "THC_CBD": {"base_size": 8, "min_size": 6, "max_length": 50},
    "RATIO": {"base_size": 8, "min_size": 6, "max_length": 30},
    "WEIGHT": {"base_size": 10, "min_size": 8, "max_length": 20},
    "UNITS": {"base_size": 10, "min_size": 8, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 7, "min_size": 1, "max_length": 40}
}


# ------------------ Helper Functions for Normalization ------------------
def normalize(val):
    return str(val).strip().lower()

def extract_float(x):
    try:
        matches = re.findall(r"[\d.]+", x)
        if matches:
            return float(matches[0])
    except Exception:
        pass
    return 0

# ------------------ UI Functions for Editing Data ------------------
def edit_data_manually():
    # Implementation for editing data manually
    file_path_val = file_entry.get()
    if not file_path_val:
        messagebox.showerror("Error", "Please upload a data file before editing.")
        return
    try:
        transformed_excel_file = preprocess_excel(file_path_val)
        logging.debug(f"Transformed file created: {transformed_excel_file}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to transform Excel file: {e}")
        return
    open_file(transformed_excel_file)
    response = messagebox.askokcancel(
        "Edit Data Manually",
        "The transformed spreadsheet has been opened in Excel.\n\n"
        "Please edit and save the file in Excel, then click OK to reload the updated data.\n"
        "If you haven't finished editing, click Cancel."
    )
    if response:
        try:
            global global_df
            global_df = pd.read_excel(transformed_excel_file, engine="openpyxl")
            populate_filter_dropdowns()
            populate_product_names()
            messagebox.showinfo("Reload Successful", "Data has been reloaded from the edited file.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to reload edited file: {e}")

# ------------------ DOCX Helper Functions ------------------
def disable_autofit(table):
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_table_cell_spacing(table, spacing_twips):
    tblPr = table._element
    tblPr_obj = tblPr.find(qn('w:tblPr'))
    if tblPr_obj is None:
        tblPr_obj = OxmlElement('w:tblPr')
        tblPr.insert(0, tblPr_obj)
    tblCellSpacing = tblPr_obj.find(qn('w:tblCellSpacing'))
    if tblCellSpacing is None:
        tblCellSpacing = OxmlElement('w:tblCellSpacing')
        tblPr_obj.append(tblCellSpacing)
    tblCellSpacing.set(qn('w:w'), str(spacing_twips))
    tblCellSpacing.set(qn('w:type'), 'dxa')

def _set_row_height_exact(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    for child in trPr.findall(qn('w:trHeight')):
        trPr.remove(child)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt.pt * 20)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cell_has_text(cell):
    return bool(cell.text.strip())

def rebuild_table_with_nonempty_cells(doc, old_table, num_cols=5):
    non_empty_texts = [
    cell.text for row in old_table.rows for cell in row.cells if cell_has_text(cell)
]

    # Create new table directly using texts:
    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    idx = 0
    for row in new_table.rows:
        for cell in row.cells:
            if idx < len(non_empty_texts):
                cell.text = non_empty_texts[idx]
                idx += 1
    old_table._element.getparent().remove(old_table._element)
    total_cells = len(non_empty_cells)
    if total_cells == 0:
        return None
    rows_needed = math.ceil(total_cells / num_cols)
    new_table = doc.add_table(rows=rows_needed, cols=num_cols)
    new_table.alignment = 1
    disable_autofit(new_table)
    tblPr = new_table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    new_table._element.insert(0, tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    fixed_col_width = "2000"
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    idx = 0
    for r in range(rows_needed):
        for c in range(num_cols):
            cell = new_table.cell(r, c)
            cell._tc.clear_content()
            if idx < total_cells:
                for child in non_empty_cells[idx]:
                    cell._tc.append(deepcopy(child))
                idx += 1
            else:
                cell.text = ""
    return new_table

# ------------------ DOCX Expand Template Functions ------------------


def expand_template_to_3x3_fixed(template_path):
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template must contain at least one table.")
    old_table = doc.tables[0]
    source_cell_xml = deepcopy(old_table.cell(0, 0)._tc)
    old_table._element.getparent().remove(old_table._element)
    # strip any leading empty paragraphs
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    new_table = doc.add_table(rows=3, cols=3)
    new_table.alignment = 1
    disable_autofit(new_table)

    # rebuild grid to fixed 3.5"x2.5" cells
    fixed_col_width = str(int(3.5 * 1440 / 3))  # total width split among 3
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    for i in range(3):
        for j in range(3):
            label_num = i * 3 + j + 1
            cell = new_table.cell(i, j)
            cell._tc.clear_content()
            new_tc = deepcopy(source_cell_xml)
            for text_el in new_tc.iter():
                if text_el.tag == qn('w:t') and text_el.text and "Label1" in text_el.text:
                    text_el.text = text_el.text.replace("Label1", f"Label{label_num}")
            cell._tc.extend(new_tc.xpath("./*"))

        # ── ADD ONLY INTERIOR CUT-GUIDELINES ──
    tblPr = new_table._element.find(qn('w:tblPr'))
    # remove any existing borders
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    # ── INSERT LIGHT-GREY BACKGROUND SHADING ──
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')    # light-grey
    tblPr.insert(0, shd)

    # ── NOW DRAW YOUR BORDERS ──
    tblBorders = OxmlElement('w:tblBorders')
    # hide outer borders
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
    # draw interior lines
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
    tblPr.append(tblBorders)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def expand_template_to_4x5_fixed_scaled(template_path, scale_factor=1.0):
    """
    Build a 4×5 grid of 2.5"×1.75" cells + cut-guidelines.
    Returns: a BytesIO buffer containing the .docx.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    # fixed grid dimensions
    num_cols, num_rows = 4, 5
    col_width_twips = str(int(2.5 * 1440))
    row_height_pts  = Pt(1.75 * 72)
    cut_line_twips  = int(0.001 * 1440)

    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    # strip leading blanks
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = 1

      # fixed layout
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')

    # ── INSERT LIGHT-GREY BACKGROUND SHADING ──
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')    # light-grey
    tblPr.insert(0, shd)

    # enforce fixed layout
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)


    # column widths
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)

    # row heights & cut-guidelines
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    borders = OxmlElement('w:tblBorders')
    for side in ('insideH','insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single"); b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3"); b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)

    # fill & rename Label1→Label20
    cnt = 1
    for r in range(num_rows):
        for c in range(num_cols):
            cell = tbl.cell(r,c)
            cell._tc.clear_content()
            tc = deepcopy(src_tc)
            for t in tc.iter(qn('w:t')):
                if t.text and 'Label1' in t.text:
                    t.text = t.text.replace('Label1', f'Label{cnt}')
            for el in tc.xpath('./*'):
                cell._tc.append(deepcopy(el))
            cnt += 1

    # add tiny spacing between cells
    from docx.oxml.shared import OxmlElement as OE
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing'); spacing.set(qn('w:w'), str(cut_line_twips)); spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



# ------------------ Autosize and Conditional Formatting ------------------
def set_run_font_size(run, font_size):
    run.font.size = font_size
    sz_val = str(int(font_size.pt * 2))
    rPr = run._element.get_or_add_rPr()
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), sz_val)

from docx.shared import Pt

def _complexity(text):
    """Combine character count and weighted word count into one score."""
    text = str(text or "")
    return len(text) + len(text.split()) * WORD_WEIGHT


def get_thresholded_font_size_by_word_count(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        # e.g. DESC on mini tags
        if comp <  30: size = 19
        elif comp <  40: size = 18
        elif comp <  50: size = 17
        elif comp < 70: size = 16
        elif comp <  90: size = 14
        elif comp < 100: size = 12
        else:           size = 10

    elif o == 'vertical':
        if comp <  30: size = 29
        elif comp <  60: size = 26
        elif comp < 100: size = 22
        elif comp < 140: size = 20
        else:           size = 18

    elif o == 'horizontal':
        if comp <  20: size = 34
        elif comp <  30: size = 32
        elif comp <  40: size = 28
        elif comp <  50: size = 26
        elif comp < 60: size = 24
        elif comp < 70: size = 22
        else:           size = 20

    else:  # fallback
        size = 14

    return Pt(size * scale_factor)


def get_thresholded_font_size_ratio(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        if comp <  20: size =  8
        elif comp <  40: size =  7
        else:           size =  6

    elif o == 'vertical':
        if comp <  20: size = 14
        elif comp <  30: size = 12
        elif comp < 100: size = 10
        elif comp < 140: size =  8
        else:           size = 10

    elif o == 'horizontal':
        if comp <  20: size = 16
        elif comp <  30: size = 14
        elif comp <  50: size = 12
        elif comp < 100: size = 10
        else:           size = 10

    else:
        size = 10

    return Pt(size * scale_factor)


def get_thresholded_font_size_brand(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        if comp <  10: size = 14
        elif comp <  30: size = 11
        elif comp <  40: size =  8
        else:           size =  7

    elif o == 'vertical':
        if comp <  20: size = 16
        elif comp <  40: size = 14
        elif comp <  80: size = 12
        else:           size = 11

    elif o == 'horizontal':
        if comp <  20: size = 18
        elif comp <  40: size = 16
        elif comp <  80: size = 12
        else:           size = 10

    return Pt(size * scale_factor)


def autosize_field_in_paragraph(para, marker_start, marker_end, font_params, orientation, font_name="Arial", bold=True, scale_factor=1.0):
    full_text = "".join(run.text for run in para.runs)
    if marker_start in full_text and marker_end in full_text:
        try:
            field_text = full_text.split(marker_start)[1].split(marker_end)[0].strip()
        except IndexError:
            return
        # Debug print:
        print(f"[DEBUG] Field text for marker {marker_start}: '{field_text}'")
        if marker_start == "PRODUCTBRAND_CENTER_START":
            new_size_val = get_thresholded_font_size_brand(field_text, orientation, scale_factor)
        elif marker_start == "DESC_START":
            new_size_val = get_thresholded_font_size_by_word_count(field_text, orientation, scale_factor)
        elif marker_start == "RATIO_START":
            new_size_val = get_thresholded_font_size_ratio(field_text, orientation, scale_factor)
        elif marker_start == "PRODUCTSTRAIN_START":
            new_size_val = Pt(font_params["base_size"])
        else:
            length = len(field_text)
            base_size = font_params["base_size"]
            max_length = font_params["max_length"]
            min_size = font_params["min_size"]
            new_size_val = Pt(base_size) if length <= max_length else Pt(max(min_size, base_size * (max_length / length)))
        # Increase size for placeholders.
        bold = True

        new_text = unescape(full_text.replace(marker_start, "").replace(marker_end, ""))
        p_element = para._element
        for child in list(p_element):
            p_element.remove(child)
        new_run = para.add_run(new_text)
        new_run.font.size = new_size_val
        new_run.font.name = font_name
        new_run.font.bold = bold
        set_run_font_size(new_run, new_size_val)
        if marker_start == "PRODUCTBRAND_CENTER_START":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def autosize_fields(doc, font_scheme, orientation, scale_factor=1.0):
    def recursive_autosize(element, marker_start, marker_end, font_params, orientation, scale_factor):
        for para in element.paragraphs:
            autosize_field_in_paragraph(para, marker_start, marker_end, font_params, orientation, scale_factor=scale_factor)
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    recursive_autosize(cell, marker_start, marker_end, font_params, orientation, scale_factor)
    recursive_autosize(doc, "DESC_START", "DESC_END", font_scheme["DESC"], orientation, scale_factor)
    recursive_autosize(doc, "PRIC_START", "PRIC_END", font_scheme["PRIC"], orientation, scale_factor)
    recursive_autosize(doc, "LINEAGE_START", "LINEAGE_END", font_scheme["LINEAGE"], orientation, scale_factor)
    recursive_autosize(doc, "LINEAGE_CENTER_START", "LINEAGE_CENTER_END", font_scheme["LINEAGE_CENTER"], orientation, scale_factor)
    recursive_autosize(doc, "THC_CBD_START", "THC_CBD_END", font_scheme["THC_CBD"], orientation, scale_factor)
    recursive_autosize(doc, "RATIO_START", "RATIO_END", font_scheme["RATIO"], orientation, scale_factor)
    recursive_autosize(doc, "WEIGHT_START", "WEIGHT_END", font_scheme["WEIGHT"], orientation, scale_factor)
    recursive_autosize(doc, "UNITS_START", "UNITS_END", font_scheme["UNITS"], orientation, scale_factor)
    recursive_autosize(doc, "PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END", font_scheme["PRODUCTSTRAIN"], orientation, scale_factor)
    recursive_autosize(doc, "PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END", font_scheme["PRODUCTBRAND_CENTER"], orientation, scale_factor)
    return doc

def apply_conditional_formatting(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_formatting_to_cell(cell)
    shade_middle_row_if_needed(doc)
    return doc

def apply_formatting_to_cell(cell):
    text = cell.text.strip().upper()

    # 1) CANNABINOID RATIOS → SHADE YELLOW
    if any(chem in text for chem in ["CBD", "CBN", "CBG", "CBC"]):
        set_cell_background(cell, "F1C232")
        set_font_color_white(cell)
        return

    # 2) PARAPHERNALIA → SHADE PINK
    if "PARAPHERNALIA" in text:
        set_cell_background(cell, "FFC0CB")
        set_font_color_white(cell)
        return

    # 3) HYBRID SUB‑TYPES
    if "HYBRID/INDICA" in text or "HYBRID INDICA" in text:
        set_cell_background(cell, "9900FF"); set_font_color_white(cell); return
    if "HYBRID/SATIVA" in text or "HYBRID SATIVA" in text:
        set_cell_background(cell, "ED4123"); set_font_color_white(cell); return

    # 4) PLAIN LINEAGES
    if "INDICA"   in text:
        set_cell_background(cell, "9900FF"); set_font_color_white(cell); return
    if "SATIVA"   in text:
        set_cell_background(cell, "ED4123"); set_font_color_white(cell); return
    if "HYBRID"   in text:
        set_cell_background(cell, "009900"); set_font_color_white(cell); return
    if "MIXED"    in text:
        set_cell_background(cell, "0021F5"); set_font_color_white(cell); return
    if "CBD"      in text:  # fallback for lone “CBD”
        set_cell_background(cell, "F1C232"); set_font_color_white(cell); return
    if "PARAPHERNALIA" in text:  # fallback catch
        set_cell_background(cell, "FFC0CB"); set_font_color_white(cell); return

    # 5) DEFAULT → white background
    set_cell_background(cell, "FFFFFF")





def set_cell_background(cell, color_hex):
    if not cell.text.strip():
        cell.text = " "
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    new_shd = OxmlElement('w:shd')
    new_shd.set(qn('w:val'), 'clear')
    new_shd.set(qn('w:color'), 'auto')
    new_shd.set(qn('w:fill'), color_hex.upper())
    tcPr.append(new_shd)

def set_font_color_white(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.highlight_color = None

def shade_middle_row_if_needed(doc, shade_color="C0C0C0"):
    for table in doc.tables:
        if not table.rows:
            continue
        mid_index = len(table.rows) // 2
        mid_row = table.rows[mid_index]
        row_text = " ".join(cell.text for cell in mid_row.cells).upper()
        if "LINEAGE" in row_text or "PRODUCT STRAIN" in row_text:
            for cell in mid_row.cells:
                set_cell_background(cell, shade_color)
                set_font_color_white(cell)
    return doc

def safe_fix_paragraph_spacing(doc):
    for para in doc.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
    return doc

def remove_extra_spacing(doc):
    try:
        normal_style = doc.styles["Normal"].paragraph_format
        normal_style.space_before = Pt(0)
        normal_style.space_after = Pt(0)
        normal_style.line_spacing = 1
    except Exception as e:
        print("Error adjusting Normal style:", e)
    return doc

def clear_cell_margins(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                for margin in ("top", "left", "bottom", "right"):
                    m = tcPr.find(qn(f"w:{margin}"))
                    if m is None:
                        m = OxmlElement(f"w:{margin}")
                        tcPr.append(m)
                    m.set(qn("w:w"), "0")
                    m.set(qn("w:type"), "dxa")
    return doc

def clear_table_cell_padding(doc):
    for table in doc.tables:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            tblCellMar = tblPr.find(qn('w:tblCellMar'))
            if tblCellMar is not None:
                for margin in ['top', 'left', 'bottom', 'right']:
                    m = tblCellMar.find(qn(f'w:{margin}'))
                    if m is not None:
                        m.set(qn('w:w'), "0")
                        m.set(qn('w:type'), "dxa")
    return doc

def compact_table_cells(doc, num_cols=3):
    if not doc.tables:
        return doc
    orig_table = doc.tables[0]
    non_blank_cells = []
    for row in orig_table.rows:
        for cell in row.cells:
            if cell.text.strip():
                non_blank_cells.append(deepcopy(cell._tc))
    orig_table._element.getparent().remove(orig_table._element)
    num_cells = len(non_blank_cells)
    num_rows = (num_cells + num_cols - 1) // num_cols
    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    new_table.alignment = 1
    fixed_col_width = "2000"
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    cell_index = 0
    for r in range(num_rows):
        for c in range(num_cols):
            cell = new_table.cell(r, c)
            cell._tc.clear_content()
            if cell_index < num_cells:
                for child in non_blank_cells[cell_index]:
                    cell._tc.append(deepcopy(child))
                cell_index += 1
            else:
                cell.text = ""
    return new_table

def reapply_table_cell_spacing_only(doc, spacing_inches=0.03):
    spacing_twips = int(spacing_inches * 1440)
    for table in doc.tables:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        tblCellSpacing = tblPr.find(qn('w:tblCellSpacing'))
        if tblCellSpacing is None:
            tblCellSpacing = OxmlElement('w:tblCellSpacing')
            tblPr.append(tblCellSpacing)
        tblCellSpacing.set(qn('w:w'), str(spacing_twips))
        tblCellSpacing.set(qn('w:type'), 'dxa')

def remove_trailing_blank_paragraphs(doc):
    """
    Remove trailing empty paragraphs from a Document to help prevent a blank page.
    """
    # Iterate over paragraphs in reverse order
    for para in reversed(doc.paragraphs):
        if not para.text.strip():  # if paragraph is blank
            # Remove the paragraph element from its parent
            p_element = para._element
            p_element.getparent().remove(p_element)
        else:
            # Stop once a non-empty paragraph is reached.
            break
    return doc

# ------------------ Excel Processing Functions ------------------
from decimal import Decimal, InvalidOperation

def format_price(p):
    try:
        value = str(p).strip().lstrip("$")
        val = float(value)
        if val.is_integer():
            return f"'{int(val)}"
        else:
            s = str(val).rstrip("0").rstrip(".")
            return f"'{s}"
    except Exception:
        return f"'{str(p).strip().lstrip('$')}"

    
def format_weight(w):
    try:
        val = float(w)
        return str(int(val)) if val.is_integer() else str(val)
    except Exception:
        return str(w)

def sanitize_filename(s):
    allowed = "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 _-&"
    return "".join(ch for ch in s if ch in allowed).replace(" ", "_")


def format_price_preprocess(p):
    try:
        s = str(p).strip()
        # Remove any leading "$"
        if s.startswith("$"):
            s = s[1:]
        # Remove apostrophes and stray whitespace
        s = s.replace("'", "").strip()
        # Convert to float
        val = float(s)
        # If the value is integer, return it without any decimal portion
        if val.is_integer():
            return f"${int(val)}"
        else:
            # Otherwise, format with 2 decimals, then remove any trailing zeros and dot
            formatted = f"{val:.2f}"
            # If formatted ends with .00, remove it completely
            if formatted.endswith(".00"):
                return f"${formatted[:-3]}"
            else:
                formatted = formatted.rstrip("0").rstrip(".")
                return f"${formatted}"
    except Exception:
        # Fallback just in case
        fallback = str(p).strip()
        if fallback.startswith("$"):
            fallback = fallback[1:]
        fallback = fallback.replace("'", "").strip()
        try:
            fv = float(fallback)
            if fv.is_integer():
                return f"${int(fv)}"
            else:
                formatted = f"{fv:.2f}"
                if formatted.endswith(".00"):
                    return f"${formatted[:-3]}"
                else:
                    return f"${formatted.rstrip('0').rstrip('.')}"
        except Exception:
            return f"${fallback}"
        
def fix_description_spacing(desc):
    """
    Inserts a space before a hyphen that is immediately followed by a digit.
    
    For example:
      "Gelato Infused Pre-Roll- 1g" becomes "Gelato Infused Pre-Roll - 1g"
      
    This function only affects hyphens directly preceding digits,
    so it leaves compound words like "Pre-Roll" intact.
    """
    # (?<!\s) ensures that the hyphen isn't already preceded by a space.
    # \s* eats up any existing whitespace after the hyphen.
    # (\d) captures the first digit that follows.
    return re.sub(r'(?<!\s)-\s*(\d)', r' - \1', desc)

def preprocess_excel(file_path, filters=None):
    import datetime, os, re
    import numpy as np
    import pandas as pd

    # 1) Read & dedupe, force-key columns to string for .str ops
    dtype_dict = {
        "Product Type*": "string",
        "Lineage": "string",
        "Product Brand": "string",
        "Vendor": "string",
        "Weight Unit* (grams/gm or ounces/oz)": "string",
        "Product Name*": "string"
    }
    df = pd.read_excel(file_path, engine="openpyxl", dtype=dtype_dict)
    df.drop_duplicates(inplace=True)

    # 2) Trim product names
    if "Product Name*" in df.columns:
        df["Product Name*"] = df["Product Name*"].str.lstrip()

    # 3) Ensure required columns exist
    for col in ["Product Type*", "Lineage", "Product Brand"]:
        if col not in df.columns:
            df[col] = "Unknown"

    # 4) Exclude sample rows
    df = df[~df["Product Type*"].isin(["Samples - Educational", "Sample - Vendor"])]

    # 5) Rename for convenience
    df.rename(columns={
        "Weight Unit* (grams/gm or ounces/oz)": "Units",
        "Price* (Tier Name for Bulk)": "Price",
        "Vendor/Supplier*": "Vendor",
        "DOH Compliant (Yes/No)": "DOH",
        "Concentrate Type": "Ratio"
    }, inplace=True)

    # 6) Normalize units
    if "Units" in df.columns:
        df["Units"] = df["Units"].str.lower().replace(
            {"ounces": "oz", "grams": "g"}, regex=True
        )

    # 7) Standardize Lineage
    if "Lineage" in df.columns:
        df["Lineage"] = (
            df["Lineage"]
              .str.lower()
              .replace({
                  "indica_hybrid": "HYBRID/INDICA",
                  "sativa_hybrid": "HYBRID/SATIVA",
                  "sativa": "SATIVA",
                  "hybrid": "HYBRID",
                  "indica": "INDICA"
              })
              .fillna("HYBRID")
              .str.upper()
        )

    # 8) Build Description & Ratio & Strain
    if "Product Name*" in df.columns:
        df["Description"] = df["Product Name*"].str.split(" by").str[0]
        mask_para = df["Product Type*"].str.strip().str.lower() == "paraphernalia"
        df.loc[mask_para, "Description"] = (
            df.loc[mask_para, "Description"]
              .str.replace(r"\s*-\s*\d+g$", "", regex=True)
        )
            # ──  REMOVE DUPLICATES BASED ON Description ───────────────
        df = df.drop_duplicates(subset=["Description"], keep="first")

           # … after this block that builds df["Ratio"] …
        df["Ratio"] = df["Product Name*"].str.extract(r"-\s*(.+)").fillna("")
        df["Ratio"] = df["Ratio"].str.replace(r" / ", " ", regex=True)

            # ── ensure “Product Strain” exists and is a Categorical ──────────────
        if "Product Strain" not in df.columns:
            df["Product Strain"] = ""
        # this will turn anything in that column into a category dtype
        df["Product Strain"] = df["Product Strain"].astype("category")

            # ‑‑‑ Force all non‑CBD‑Blend strains to "Mixed" ───────────────────────
        if "Product Strain" in df.columns:
            # Convert to plain string then override
            df["Product Strain"] = df["Product Strain"].astype(str).apply(
                lambda s: "CBD Blend" if s == "CBD Blend" else "Mixed"
            ).astype("category")


        # ── now force CBD Blend for any ratio containing CBD, CBC, CBN or CBG ──
        mask_cbd_ratio = df["Ratio"].str.contains(
            r"\b(?:CBD|CBC|CBN|CBG)\b", case=False, na=False
        )
        if mask_cbd_ratio.any():
            # add “CBD Blend” to the categories if it’s not already there
            if "CBD Blend" not in df["Product Strain"].cat.categories:
                df["Product Strain"] = df["Product Strain"].cat.add_categories(["CBD Blend"])
            # assign
            df.loc[mask_cbd_ratio, "Product Strain"] = "CBD Blend"



        # ── SPECIAL CASE: anything with Product Type “paraphernalia” gets
    #    its Product Strain forcibly set to "Paraphernalia"
    mask_para = df["Product Type*"].str.strip().str.lower() == "paraphernalia"

    # ensure the column exists as categorical and add the new category
    if "Product Strain" not in df.columns:
        df["Product Strain"] = pd.Categorical([], categories=["Paraphernalia"])
    else:
        # if it’s already categorical, just add the new category
        if isinstance(df["Product Strain"].dtype, pd.CategoricalDtype):
            if "Paraphernalia" not in df["Product Strain"].cat.categories:
                df["Product Strain"] = df["Product Strain"].cat.add_categories(["Paraphernalia"])
        else:
            # not categorical yet → make it categorical with this extra
            df["Product Strain"] = pd.Categorical(df["Product Strain"], 
                                                  categories=list(df["Product Strain"].unique()) + ["Paraphernalia"])

    # now you can safely assign
    df.loc[mask_para, "Product Strain"] = "Paraphernalia"


    # 9) Convert key fields to categorical
    for col in ["Product Type*", "Lineage", "Product Brand", "Vendor"]:
        if col in df.columns:
            df[col] = df[col].astype("category")

    # 10) CBD overrides
    if "Description" in df.columns and "Lineage" in df.columns:
        cbd_mask = df["Description"].str.contains(
            r"CBD|CBN|CBC|CBG|:", case=False, na=False
        )
        if "CBD" not in df["Lineage"].cat.categories:
            df["Lineage"] = df["Lineage"].cat.add_categories(["CBD"])
        df.loc[cbd_mask, "Lineage"] = "CBD"
    if "Description" in df.columns and "Product Strain" in df.columns:
        cbd_mask = df["Description"].str.contains(
            r"CBD|CBN|CBC|CBG|:", case=False, na=False
        )
        if "CBD Blend" not in df["Product Strain"].cat.categories:
            df["Product Strain"] = df["Product Strain"].cat.add_categories(["CBD Blend"])
        df.loc[cbd_mask, "Product Strain"] = "CBD Blend"

    # 11) Trim any extra columns
    if df.shape[1] > 41:
        df = df.iloc[:, :41]

    # 12) Normalize Weight* and CombinedWeight
    if "Weight*" in df.columns:
        df["Weight*"] = pd.to_numeric(df["Weight*"], errors="coerce") \
            .apply(lambda x: str(int(x)) if pd.notnull(x) and float(x).is_integer() else str(x))
    if "Weight*" in df.columns and "Units" in df.columns:
        df["CombinedWeight"] = (df["Weight*"] + df["Units"]).astype("category")

    # 13) Format Price
    if "Price" in df.columns:
        def format_p(p):
            s = str(p).strip().lstrip("$").replace("'", "").strip()
            try:
                v = float(s)
                return f"${int(v)}" if v.is_integer() else f"${v:.2f}"
            except:
                return f"${s}"
        df["Price"] = df["Price"].apply(lambda x: format_p(x) if pd.notnull(x) else "")
        df["Price"] = df["Price"].astype("string")

    # 14) Special pre-roll Ratio logic
    def process_ratio(row):
        t = str(row.get("Product Type*", "")).strip().lower()
        if t in ["pre-roll", "infused pre-roll"]:
            parts = str(row.get("Ratio", "")).split(" - ")
            if len(parts) >= 3:
                new = " - ".join(parts[2:]).strip()
            elif len(parts) == 2:
                new = parts[1].strip()
            else:
                new = parts[0].strip()
            return f" - {new}" if not new.startswith(" - ") else new
        return row.get("Ratio", "")
        # … SPECIAL pre-roll Ratio logic ────────────────────────────────
    df["Ratio"] = df.apply(process_ratio, axis=1)

    # … (suffix-building and Excel output) …
    today = datetime.datetime.today().strftime("%Y-%m-%d")
    suffix = "all"  # or built from `filters`
    out = os.path.join(
        os.path.expanduser("~"),
        "Downloads",
        f"{today}_{suffix}.xlsx"
    )
    df.to_excel(out, index=False, engine="openpyxl")
    return out


def chunk_records(records, chunk_size=4):
    """Yield successive n‑sized chunks from the list of records."""
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

def no_filters_selected():
    filters = [
        product_type_filter_var.get(),
        lineage_filter_var.get(),
        product_brand_filter_var.get(),
        vendor_filter_var.get(),
        weight_filter_var.get(),
        product_strain_filter_var.get()
    ]
    return all(f == "All" for f in filters)



# ------------------ Processing Functions ------------------
def process_chunk(args):
    """
    Processes a chunk of records and returns a DOCX document as a bytes buffer.
    """
    from io import BytesIO
    from docx import Document
    from docxtpl import DocxTemplate
    from docx.shared import Mm
    # unpack
    chunk, base_template, font_scheme, orientation, scale_factor = args

    # prepare template buffer
    if orientation == "mini":
        local_template_buffer = expand_template_to_4x5_fixed_scaled(
            base_template,
            scale_factor=scale_factor
        )
    else:
        # for horizontal/vertical/inventory you would use a different expand function
        local_template_buffer = base_template  # or appropriate buffer

    tpl = DocxTemplate(local_template_buffer)

    # build context and render…
    context = {}

    image_width = Mm(8) if orientation == "mini" else Mm(12 if orientation == 'vertical' else 14)
    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
    
    if orientation == "mini":
        num_labels = 25
    elif orientation == "inventory":
        num_labels = 4
    else:
        num_labels = 9

    for i in range(num_labels):
        label_data = {}
        if i < len(chunk):
            row = chunk[i]
            doh_value = str(row.get("DOH", "")).strip()
            product_type = str(row.get("Product Type*", "")).strip().lower()
            if doh_value == "Yes":
                high_cbd_types = [
                    "high cbd edible liquid - doh compliant",
                    "high cbd edible solid - doh compliant",
                    "high cbd topical - doh compliant"
                ]
                if product_type in high_cbd_types:
                    high_cbd_image_path = resource_path(os.path.join("templates", "HighCBD.png"))
                    label_data["DOH"] = InlineImage(tpl, high_cbd_image_path, width=image_width)
                else:
                    label_data["DOH"] = InlineImage(tpl, doh_image_path, width=image_width)
            else:
                label_data["DOH"] = ""
                
            price_val = f"{row.get('Price', '')}"
            label_data["Price"] = wrap_with_marker(price_val, "PRIC")
            
            lineage_text   = str(row.get("Lineage", "")).strip()
            product_brand  = str(row.get("Product Brand", "")).strip()
            label_data["ProductBrand"] = wrap_with_marker(product_brand.upper(), "PRODUCTBRAND_CENTER")

            # ── SPECIAL CASE: paraphernalia shows Vendor instead of Brand ──
          # ── SPECIAL CASE: paraphernalia ──────────────────────────────────
            if orientation not in ["mini", "inventory"] and product_type == "paraphernalia":
                vendor_text = str(row.get("Vendor", "")).strip()
                # show vendor in the Lineage cell
                label_data["Lineage"]         = wrap_with_marker(vendor_text.upper(), "PRODUCTBRAND_CENTER")
                # no THC/CBD block
                label_data["Ratio_or_THC_CBD"] = ""
                # force Product Strain to read "Paraphernalia"
                label_data["ProductStrain"]    = wrap_with_marker("Paraphernalia", "PRODUCTSTRAIN")
                # remove any weight/units field
                label_data["WeightUnits"]      = ""


            # ── all other types unchanged ─────────────────────────────────
            elif orientation not in ["mini", "inventory"]:
                # these two extract types also get a THC/CBD block
                if product_type in {"co2 concentrate","alcohol/ethanol extract"}:
                    label_data["Lineage"]          = wrap_with_marker(lineage_text, "LINEAGE")
                    label_data["Ratio_or_THC_CBD"] = wrap_with_marker("THC:\n\nCBD:", "THC_CBD")
                    label_data["ProductStrain"]    = ""
                elif product_type in {"flower", "vape cartridge", "solventless concentrate",
                                    "concentrate", "pre-roll", "infused pre-roll"}:
                    label_data["Lineage"]          = wrap_with_marker(lineage_text, "LINEAGE")
                    label_data["Ratio_or_THC_CBD"] = wrap_with_marker("THC:\n\nCBD:", "THC_CBD")
                    label_data["ProductStrain"]    = ""
                else:
                    label_data["Lineage"]          = wrap_with_marker(product_brand.upper(), "PRODUCTBRAND_CENTER")
                    label_data["Ratio_or_THC_CBD"] = wrap_with_marker(row.get("Ratio", ""), "RATIO")
                    label_data["ProductStrain"]    = wrap_with_marker(row.get("Product Strain", ""), "PRODUCTSTRAIN")

            
            label_data["ProductBrandFontSize"] = get_thresholded_font_size_brand(product_brand, scale_factor=1.0)
            
            def format_ratio_multiline(ratio_text):
                if not isinstance(ratio_text, str):
                    return ""
                parts = re.split(r"\s*\|\s*|\s{2,}", ratio_text.strip())
                return "\n".join(p.strip() for p in parts if p.strip())
            import re   # at top of your MAIN.py

            # … after you’ve pulled `product_type = str(row.get("Product Type*", "")).lower()` …

            # coerce to string, strip leading/trailing whitespace
            raw_ratio = row.get("Ratio", "") or ""
            # force to str and strip *all* leading/trailing whitespace
            clean_ratio = str(raw_ratio).strip()
            label_data["Ratio"] = wrap_with_marker(
                format_ratio_multiline(clean_ratio),
                "RATIO"
)


            # clean up the description cell:
            raw_desc = str(row.get("Description", "")).strip()
            if product_type == "paraphernalia":
                cleaned_desc = re.sub(r"\s*-\s*\d+g$", "", raw_desc)
            else:
                cleaned_desc = raw_desc

            # ── THIS IS WHERE WE ADD THE TRAILING NBSP FOR PRE-ROLLS ──────────────
            if product_type in {"pre-roll", "infused pre-roll"}:
                cleaned_desc += "\u00A0"

            label_data["Description"] = wrap_with_marker(cleaned_desc, "DESC")


                

            label_data["Description"] = wrap_with_marker(cleaned_desc, "DESC")

            
                       # … after you’ve extracted raw_desc …
            # get the numeric weight and its unit
            try:
                weight_val = float(row.get("Weight*", ""))
            except Exception:
                weight_val = None
            units_val = row.get("Units", "").lower()    # e.g. 'g' or 'oz'

            # ── NEW: convert certain gram‑based products to oz ─────────────
            edible_types = {
                "edible (solid)",
                "edible (liquid)",
                "high cbd edible liquid",   # treat this the same as “edible (liquid)”
                "tincture",
                "topical",
                "capsule",
}

            if product_type in edible_types and units_val in {"g", "grams"} and weight_val is not None:
                weight_val = weight_val * 0.03527396195
                units_val = "oz"

            # now build the display string
            if weight_val is not None and units_val:
                weight_str = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                weight_units = ""

            # … earlier in process_chunk, after you have:
            # … after reading weight_val, units_val …
            product_type = product_type.lower()
            units_val    = units_val.lower()

            # include both “edible (liquid)” and “high cbd edible liquid”
            edible_types = {
                "edible (solid)",
                "edible (liquid)",
                "high cbd edible liquid",
                "tincture",
                "topical",
                "capsule",
            }

            # if it's one of our edible types stored in grams, convert to oz
            if product_type in edible_types and units_val in {"g", "grams"} and weight_val is not None:
                weight_val *= 0.03527396195
                units_val = "oz"

            # now build your display string as before
            if weight_val is not None and units_val:
                weight_str   = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                weight_units = ""

            # Compute the normal weight string:
            # Compute the normal weight string:
            try:
                weight_val = float(row.get("Weight*", ""))
            except:
                weight_val = None
            units_val = row.get("Units", "")
            if weight_val is not None and units_val:
                weight_str = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                normal_weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                normal_weight_units = ""

            # SPECIAL OVERRIDE FOR PRE-ROLL & INFUSED PRE-ROLL
            # — final override for WeightUnits —
            if product_type in {"pre-roll", "infused pre-roll"}:
                # for pre-rolls, shove the Ratio into the WeightUnits slot
                formatted = format_ratio_multiline(row.get("Ratio", ""))
                label_data["WeightUnits"] = wrap_with_marker(formatted, "RATIO")
            elif product_type == "paraphernalia":
                # hide weight for paraphernalia
                label_data["WeightUnits"] = ""
            elif orientation == "inventory":
                # inventory slips show actual weight_units (from earlier)
                label_data["WeightUnits"] = weight_units
            else:
                # everyone else gets the normal weight+unit
                label_data["WeightUnits"] = normal_weight_units
            # final override for WeightUnits
            if orientation == "inventory":
                # keep the usual weight on inventory slips
                label_data["WeightUnits"] = weight_units

            elif product_type in {"pre-roll", "infused pre-roll"}:
                raw_ratio = row.get("Ratio", "")
                label_data["WeightUnits"] = wrap_with_marker(
                    format_ratio_multiline(raw_ratio),
                    "DESC"
                )

            elif product_type == "paraphernalia":
                # hide weight completely on paraphernalia labels
                label_data["WeightUnits"] = ""

            else:
                # everyone else gets the normal weight+unit
                label_data["WeightUnits"] = normal_weight_units

            
        else:
            label_data = {
                "DOH": "",
                "Price": "",
                "Lineage": "",
                "Ratio_or_THC_CBD": "",
                "ProductBrand": "",
                "Ratio": "",
                "Description": "",
                "ProductStrain": "",
                "WeightUnits": ""
            }
            if orientation == "inventory":
                label_data["AcceptedDate"] = ""
                label_data["Vendor"] = ""
                label_data["Barcode"] = ""
                label_data["ProductName"] = ""
                label_data["ProductType"] = ""
                label_data["QuantityReceived"] = ""
        context[f"Label{i+1}"] = label_data

    tpl.render(context)

    # save into buffer
    buffer = BytesIO()
    tpl.docx.save(buffer)
    buffer.seek(0)

    # load into python-docx for downstream fixes
    doc = Document(buffer)

    if orientation != "inventory":
        autosize_fields(doc, font_scheme, orientation, scale_factor=scale_factor)
        apply_conditional_formatting(doc)
        safe_fix_paragraph_spacing(doc)
        remove_extra_spacing(doc)
        clear_cell_margins(doc)
        clear_table_cell_padding(doc)

    if orientation == "mini":
        # clear truly empty cells on mini sheets
        def fully_clear_cell(cell):
            for child in list(cell._tc):
                cell._tc.remove(child)
            set_cell_background(cell, "FFFFFF")
            cell.text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        fully_clear_cell(cell)

    # return raw bytes
    final_buffer = BytesIO()
    doc.save(final_buffer)
    return final_buffer.getvalue()

# ------------------ Run Full Process Functions ------------------
def filter_column(df, column, var):
    filter_val = normalize(var.get())
    if filter_val != "all" and column in df.columns:
        return df[df[column].astype(str).apply(normalize) == filter_val]
    return df

from io import BytesIO
from docx import Document
from docxcompose.composer import Composer
import datetime
import os
import pandas as pd
from tkinter import messagebox

from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ROW_HEIGHT_RULE

from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared    import Inches

def process_name_chunk(args):
    """
    Given a chunk of records, a 3×3 template buffer, and the orientation,
    build a back‐side page with product names.  For vertical tags we
    swap to 2.5"×3.5" cells.
    """
    chunk, template_buffer, orientation = args

    buf = BytesIO(template_buffer.getvalue())
    doc = Document(buf)

    table = doc.tables[0]
    disable_autofit(table)

    # choose cell dimensions by orientation
    if orientation == "vertical":
        col_w = Inches(2.3)
        row_h = Inches(3.3)
    else:
        col_w = Inches(3.3)
        row_h = Inches(2.3)

    # apply widths/heights
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_w

    for row in table.rows:
        row.height = row_h
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # fill in names (or leave blank)
    total = len(table.columns) * len(table.rows)
    for idx in range(total):
        r, c = divmod(idx, len(table.columns))
        cell = table.cell(r, c)
        if idx < len(chunk):
            cell.text = chunk[idx]["Product Name*"]
        else:
            cell.text = ""

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def add_vendor_back_pages(master_doc, records):
    """
    After each front chunk (3×3) in master_doc, append a back page
    with the same exact table dimensions, populated with Vendor/Brand.
    """
    # first, grab the already-expanded 3×3 grid from the front
    # assume you expanded once with expand_template_to_3x3_fixed()
    # and that you passed that buffer in as `fixed_buf`
    # so let’s stash it on your master_doc for reuse:
    fixed_buf = master_doc._fixed_3x3_buffer

    # iterate your record‐chunks of 9
    for chunk in chunk_records(records, chunk_size=9):
        # load a fresh Document from the same buffer
        back_doc = Document(BytesIO(fixed_buf.getvalue()))
        tbl      = back_doc.tables[0]

        # ensure each row/col is exactly 3.5"×2.5"
        # (these are the same numbers your expand_template function used)
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = Inches(3.5)
        for row in tbl.rows:
            row.height = Inches(2.5)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # now fill with Vendor / Brand
        for idx, rec in enumerate(chunk):
            r, c = divmod(idx, 3)
            cell = tbl.cell(r, c)
            vendor = rec.get("Vendor", "").strip()
            brand  = rec.get("Product Brand", "").strip()
            cell.text = f"{vendor}\n{brand}" if vendor or brand else ""

        # append directly after the last section of master_doc
        composer = Composer(master_doc)
        composer.append(back_doc)


# ─── Main generation function ─────────────────────────────────────
def run_full_process_by_group(template_type, group_by_fields=["Lineage", "Product Strain"]):
    import io, os, datetime
    from docx import Document
    from docxcompose.composer import Composer
    splash = show_splash2(root)

    # 1) Get & validate file
    file_path = file_entry.get()
    if not file_path:
        messagebox.showerror("Error", "Please select a data file.")
        return

    # 2) Preprocess & reload
    filters = {
        "product_type": product_type_filter_var.get(),
        "lineage":      lineage_filter_var.get(),
        "brand":        product_brand_filter_var.get(),
        "vendor":       vendor_filter_var.get(),
        "weight":       weight_filter_var.get(),
        "strain":       product_strain_filter_var.get()
    }
    
    prepped = preprocess_excel(file_path, filters)
    global global_df
    global_df = pd.read_excel(prepped, engine="openpyxl")
    df = global_df.copy()

    # 3) Apply dropdown filters
    df = filter_column(df, "Product Type*", product_type_filter_var)
    df = filter_column(df, "Lineage",         lineage_filter_var)
    df = filter_column(df, "Product Brand",   product_brand_filter_var)
    df = filter_column(df, "Vendor",          vendor_filter_var)
    df = filter_column(df, "CombinedWeight",  weight_filter_var)
    df = filter_column(df, "Product Strain",  product_strain_filter_var)
    
    LINEAGE_ORDER = [
        "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA",
        "HYBRID/INDICA", "CBD", "MIXED", "PARAPHERNALIA"
    ]

    # assign an ordering key
    df["_lin_order"] = df["Lineage"].apply(
        lambda x: LINEAGE_ORDER.index(x) if x in LINEAGE_ORDER else len(LINEAGE_ORDER)
    )
    # secondary sort by name (optional)
    df = df.sort_values(by=["_lin_order", "Product Name*"])
    # drop the helper column
    df = df.drop(columns=["_lin_order"])

    records = df.to_dict(orient="records")


    # 4) Limit to checked Selected Tags
    sel = [n for n,v in selected_tags_vars.items() if v.get()]
    if not sel:
        messagebox.showerror("Error", "No selected tags are checked!")
        return
    df = df[df["Product Name*"].isin(sel)]

    # 5) Choose template & scheme
    if template_type == "horizontal":
        tpl_path    = resource_path("templates/horizontal.docx")
        orientation = "horizontal"
        scheme      = FONT_SCHEME_HORIZONTAL
    else:
        tpl_path    = resource_path("templates/vertical.docx")
        orientation = "vertical"
        scheme      = FONT_SCHEME_VERTICAL

    # 6) Expand to fixed 3×3
    # 6) Expand the 3×3 template once
    fixed_buf = expand_template_to_3x3_fixed(tpl_path)

    # 7) Render front & back for each chunk
    records = df.to_dict("records")
    # inside run_full_process_by_group(...)
    bytes_list = []
    for chunk in chunk_records(records, chunk_size=9):
        # front side
        front_bytes = process_chunk((chunk, fixed_buf, scheme, orientation, SCALE_FACTOR))
        bytes_list.append(front_bytes)

        # back side (same grid but sized per orientation)
        if print_vendor_back_var.get():
            back_bytes = process_name_chunk((chunk, fixed_buf, orientation))
            bytes_list.append(back_bytes)


    if not bytes_list:
        messagebox.showerror("Error", "No documents generated.")
        return

    # 8) Stitch into a master_doc
    master_doc = Document(io.BytesIO(bytes_list[0]))
    composer   = Composer(master_doc)
    for b in bytes_list[1:]:
        composer.append(Document(io.BytesIO(b)))

    # 9) Final spacing
    reapply_table_cell_spacing_only(master_doc, spacing_inches=0.03)

    # 10) Save & open
    today = datetime.datetime.now().strftime("%Y%m%d")
    suffix = "_".join(
        p for p in [filters["product_type"], filters["lineage"],
                    filters["brand"], filters["vendor"],
                    filters["weight"], filters["strain"]]
        if p and p!="All"
    ) or "all"
    out = os.path.join(os.path.expanduser("~"), "Downloads",
                       f"{today}_{orientation}_{suffix}_tags.docx")
    master_doc.save(out)
    open_file(out)
    splash.destroy()
    messagebox.showinfo("Success", f"Saved: {out}")


def run_full_process_mini(bypass_tag_filter: bool = False):
    splash = show_splash2(root)
    # ── 0.  Mini template + constants ───────────────────────────────
    base_template       = resource_path("templates/mini.docx")
    orientation         = "mini"
    current_font_scheme = FONT_SCHEME_MINI

    # ── 1.  Pick up the user’s file and filters ─────────────────────
    file_path_val = file_entry.get()
    if not file_path_val:
        messagebox.showerror("Error", "Please select a data file.")
        return

    filters = {
        "product_type": product_type_filter_var.get(),
        "lineage":      lineage_filter_var.get(),
        "brand":        product_brand_filter_var.get(),
        "vendor":       vendor_filter_var.get(),
        "weight":       weight_filter_var.get(),
        "strain":       product_strain_filter_var.get()
    }

    # preprocess once – returns XLSX path + DataFrame cached in RAM
    new_excel_file = preprocess_excel(file_path_val, filters)
    global global_df
    global_df = pd.read_excel(new_excel_file, engine="openpyxl")
    df = global_df.copy()

    # ── 2.  Apply dropdown filters & selected‑tag filter ────────────
    df = filter_column(df, "Product Type*", product_type_filter_var)
    df = filter_column(df, "Lineage",         lineage_filter_var)
    df = filter_column(df, "Product Brand",   product_brand_filter_var)
    df = filter_column(df, "Vendor",          vendor_filter_var)
    df = filter_column(df, "CombinedWeight",  weight_filter_var)
    df = filter_column(df, "Product Strain",  product_strain_filter_var)

    if "Price" in df.columns:
        df["Price"] = df["Price"].apply(lambda x: x.lstrip("'") if isinstance(x, str) else x)

    if not bypass_tag_filter:
        selected_names = [n for n, v in selected_tags_vars.items() if v.get()]
        if not selected_names:
            messagebox.showerror("Error", "No selected tags are checked!")
            return
        df = df[df["Product Name*"].isin(selected_names)]

    if df.empty:
        messagebox.showerror("Error", "No records found after filtering.")
        return

    # ── 3.  Build work items for the pool ───────────────────────────
    records   = df.to_dict(orient="records")
    base_buf = expand_template_to_4x5_fixed_scaled(
        base_template,
        scale_factor=SCALE_FACTOR
    )


    def chunk_records_mini(rec, size=30):      # bigger chunks = faster
        for i in range(0, len(rec), size):
            yield rec[i:i+size]

    work_items = [
        (chunk, base_buf, current_font_scheme, orientation, SCALE_FACTOR)
        for chunk in chunk_records_mini(records)
    ]

    # ── 4.  Render in parallel ──────────────────────────────────────
    from concurrent.futures import ProcessPoolExecutor
    with ProcessPoolExecutor(max_workers=os.cpu_count()) as exe:
        docs_bytes = list(exe.map(process_chunk, work_items))

    docs = [Document(BytesIO(b)) for b in docs_bytes if b]
    if not docs:
        messagebox.showerror("Error", "No documents were generated.")
        return

    # ── 5.  Stitch docs and save ────────────────────────────────────
        # after you collect docs = [Document(BytesIO(b)) for b in docs_bytes if b]
    blank_doc = Document()  # completely empty document for a blank “back-side” page

    master_doc = docs[0]
    composer   = Composer(master_doc)
    # start at 1 so we don’t prepend a blank in front of the very first
    for sub_doc in docs[1:]:
        composer.append(sub_doc)
        composer.append(blank_doc)   # <-- insert a blank page after each mini page

    # now continue with reapply_table_cell_spacing_only, save, open, etc.

    
    reapply_table_cell_spacing_only(master_doc)
        # after you generate the front‐side pages in master_doc…
    #if print_vendor_back_var.get():
        # your existing code that inserts the matching
        # vendor‐name back‐pages in the same grid
        #add_vendor_back_pages(master_doc, records)

    today = datetime.datetime.today().strftime("%Y-%m-%d")
    safe = lambda v: str(v).replace(" ", "").replace("/", "").replace("-", "").replace("*", "") if v and v != "All" else None
    suffix_parts = [safe(product_type_filter_var.get()),
                    safe(lineage_filter_var.get()),
                    safe(product_brand_filter_var.get()),
                    safe(vendor_filter_var.get()),
                    safe(weight_filter_var.get()),
                    safe(product_strain_filter_var.get())]
    suffix = "_".join(p for p in suffix_parts if p) or "all"

    doc_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{today}_mini_{suffix}_tags.docx")
    master_doc.save(doc_path)
    splash.destroy()
    open_file(doc_path)
    messagebox.showinfo("Success", f"Word file saved as:\n{doc_path}")



def chunk_records_inv(records, chunk_size=4):
    """Yield chunks of records where each chunk is sized for inventory slip (4 records per slip)."""
    for i in range(0, len(records), chunk_size):
        yield records[i:i+chunk_size]


def export_data_only():
    messagebox.showinfo("Export Data", "Exported data successfully.")


# ------------------ Global Mousewheel Handler ------------------
def global_mousewheel_handler(event):
    """This handler, bound with bind_all, scrolls the active canvas."""
    global current_canvas
    if current_canvas is None:
        return
    system = platform.system()
    if system == "Darwin":
        # Increase factor as needed for macOS trackpads (try 10 if necessary)
        factor = 10
        scroll_units = int(event.delta * factor)
    else:
        scroll_units = int(event.delta / 120)
    current_canvas.yview_scroll(-scroll_units, "units")
    return "break"

# Bind the global mousewheel handler at the root level.
def bind_global_mousewheel(root):
    root.bind_all("<MouseWheel>", global_mousewheel_handler)
    # For Linux:
    root.bind_all("<Button-4>", global_mousewheel_handler)
    root.bind_all("<Button-5>", global_mousewheel_handler)


# ------------------ UI Helper Functions ------------------
dropdown_cache = {}

def on_mousewheel(event, canvas):
    system = platform.system()
    if system == "Darwin":
        # On macOS, trackpad delta values are small;
        # multiply them by a factor (adjust factor as needed)
        factor = 5  # Experiment with this value.
        scroll_units = int(event.delta * factor)
    else:
        scroll_units = int(event.delta / 120)
    canvas.yview_scroll(-scroll_units, "units")
    return "break"

def update_available_tags_all_state_available():
    # Loop through each available tag's BooleanVar and set it to the checkbox state.
    for tag, var in available_tags_vars.items():
        var.set(available_tags_all_var.get())

def select_all_available():
    for var in available_tags_vars.values():
        var.set(True)
        
def build_dropdown_cache(df):
    global dropdown_cache
    cols = ["Product Type*", "Lineage", "Product Brand", "Vendor", "CombinedWeight", "Product Strain"]
    for col in cols:
        if col in df.columns:
            unique_vals = sorted(df[col].dropna().unique().tolist())
            dropdown_cache[col] = unique_vals

def update_option_menu(option_widget, var, colname):
    menu = option_widget["menu"]
    menu.delete(0, "end")
    options = dropdown_cache.get(colname, []).copy()
    options.insert(0, "All")
    for val in options:
        menu.add_command(label=val, command=lambda v=val: var.set(v))

def populate_filter_dropdowns():
    global global_df
    if global_df is None:
        return
    build_dropdown_cache(global_df)
    if "Product Type*" in global_df.columns:
        update_option_menu(product_type_option, product_type_filter_var, "Product Type*")
    if "Lineage" in global_df.columns:
        update_option_menu(lineage_option, lineage_filter_var, "Lineage")
    if "Product Brand" in global_df.columns:
        update_option_menu(product_brand_option, product_brand_filter_var, "Product Brand")
    if "Vendor" in global_df.columns:
        update_option_menu(vendor_option, vendor_filter_var, "Vendor")
    if "CombinedWeight" in global_df.columns:
        update_option_menu(weight_option, weight_filter_var, "CombinedWeight")
    if "Product Strain" in global_df.columns:
        update_option_menu(product_strain_option, product_strain_filter_var, "Product Strain")

def update_all_dropdowns():
    global _UPDATING_FILTERS, global_df, json_matched_names
    if _UPDATING_FILTERS:
        return
    _UPDATING_FILTERS = True
    try:
        # 1) Start from the full sheet
        df = global_df.copy()

        # 2) Apply each dropdown filter to df
        def apply(col, var):
            v = normalize(var.get())
            if v and v != "all" and col in df:
                return df[df[col].astype(str).apply(normalize) == v]
            return df

        df = apply("Product Type*",    product_type_filter_var)
        df = apply("Lineage",           lineage_filter_var)
        df = apply("Product Brand",     product_brand_filter_var)
        df = apply("Vendor",            vendor_filter_var)
        df = apply("Product Strain",    product_strain_filter_var)
        df = apply("CombinedWeight",    weight_filter_var)

        # 3) Rebuild every dropdown’s menu from your cached universe
        _update_option_menu(product_type_option,   product_type_filter_var,   "Product Type*",  dropdown_cache["Product Type*"])
        _update_option_menu(lineage_option,        lineage_filter_var,         "Lineage",         dropdown_cache["Lineage"])
        _update_option_menu(product_brand_option,  product_brand_filter_var,   "Product Brand",   dropdown_cache["Product Brand"])
        _update_option_menu(vendor_option,         vendor_filter_var,          "Vendor",          dropdown_cache["Vendor"])
        _update_option_menu(product_strain_option, product_strain_filter_var,  "Product Strain",  dropdown_cache["Product Strain"])
        _update_option_menu(weight_option,         weight_filter_var,          "CombinedWeight",  sorted(
            df["CombinedWeight"].dropna().unique(),
            key=lambda x: extract_float(str(x))
        ))

        # 4) Decide which list of names to show
        if json_matched_names:
            names = json_matched_names
        else:
            # now *use* the already-filtered df
            names = sorted(df["Product Name*"].dropna().unique())

        # 5) Exactly one redraw of the Available Tags panel
        populate_available_tags(names)

    finally:
        _UPDATING_FILTERS = False



def _update_option_menu(menu_widget, var, colname, value_list):
    """
    Clears and repopulates the OptionMenu.
    Always ensures 'All' is first and preserves the current selection if still valid.
    """
    menu = menu_widget["menu"]
    menu.delete(0, "end")

    all_vals = ["All"] + list(value_list)
    current = var.get()
    if current not in all_vals:
        current = "All"
    var.set(current)

    for v in all_vals:
        menu.add_command(label=v, command=lambda _v=v: var.set(_v))


def populate_available_tags(names):
    """
    Populate the left-hand ‘Available Tag List’ with colored checkbuttons,
    matching the same LINEAGE_COLOR_MAP logic as for selected tags.
    """
    global available_tags_container, available_tags_vars, placeholder_img, global_df, available_canvas

    # clear out old widgets
    for widget in available_tags_container.winfo_children():
        widget.destroy()
    available_tags_vars.clear()


    CLASSIC_TYPES = {
    "flower", "pre-roll", "concentrate",
    "infused pre-roll", "solventless concentrate",
    "vape cartridge"
    }

    for name in names:
        # get the row
        row = global_df[global_df["Product Name*"] == name].iloc[0]
        ptype = str(row["Product Type*"]).strip().lower()

        # 1) if it’s a classic type, color by Lineage
        if ptype in CLASSIC_TYPES:
            lin = str(row["Lineage"]).upper()
        else:
            # non-classic: color by strain overrides
            if ptype == "paraphernalia":
                lin = "PARAPHERNALIA"
            elif str(row["Product Strain"]) == "CBD Blend":
                lin = "CBD"
            elif str(row["Product Strain"]) == "Mixed":
                lin = "MIXED"
            else:
                # fallback to Lineage if strain isn’t one of above
                lin = str(row["Lineage"]).upper()

        # now pick your colors
        bg = LINEAGE_COLOR_MAP.get(lin, "#FFFFFF")
        fg = "white" if bg != "#FFFFFF" else "black"

        # build the row frame + transparent Checkbutton as before
        frame = tkmod.Frame(available_tags_container, bg=bg)
        frame.pack(fill="x", pady=1)

        var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(
            frame, text=name, variable=var,
            bg=bg, fg=fg, selectcolor=bg,
            activebackground=bg, activeforeground=fg,
            anchor="w", bd=0, highlightthickness=0
        )
        chk.pack(fill="x", padx=5, pady=2)
        available_tags_vars[name] = var
    available_canvas.update_idletasks()
    available_canvas.configure(scrollregion=available_canvas.bbox("all"))
    available_canvas.yview_moveto(0)

SELECTED_GROUP_ORDER = [
    "SATIVA",
    "HYBRID/SATIVA",
    "INDICA",
    "HYBRID/INDICA",
    "HYBRID",        # catch-all hybrid
    "CBD",
    "MIXED",
    "PARAPHERNALIA"
]

def _selected_lin_group(name):
    """
    Determine the lineage group for sorting/coloring,
    matching the logic in populate_selected_tags.
    """
    row = global_df[global_df["Product Name*"] == name].iloc[0]
    ptype = str(row["Product Type*"]).strip().lower()

    # classic types color by lineage
    CLASSIC_TYPES = {
        "flower", "pre-roll", "concentrate",
        "infused pre-roll", "solventless concentrate",
        "vape cartridge"
    }

    if ptype in CLASSIC_TYPES:
        lin = str(row["Lineage"]).upper()
    else:
        # paraphernalia always pink
        if ptype == "paraphernalia":
            lin = "PARAPHERNALIA"
        # CBD Blend → yellow
        elif str(row["Product Strain"]) == "CBD Blend":
            lin = "CBD"
        # everything else forced Mixed
        elif str(row["Product Strain"]) == "Mixed":
            lin = "MIXED"
        else:
            lin = str(row["Lineage"]).upper()
    # safe‐guard
    return SELECTED_GROUP_ORDER.index(lin) if lin in SELECTED_GROUP_ORDER else len(SELECTED_GROUP_ORDER)


def populate_selected_tags(names):
    splash = show_splash2(root)
    global selected_tags_container, selected_tags_vars, placeholder_img, global_df, selected_canvas

# clear out old widgets
    for widget in selected_tags_container.winfo_children():
        widget.destroy()
    selected_tags_vars.clear()

    # define the exact same lineage order you use in LINEAGE_COLOR_MAP
    lineage_buckets = [
        "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA",
        "HYBRID/INDICA", "CBD", "MIXED", "PARAPHERNALIA"
    ]

    # build a dict: lineage → [product names]
    buckets = {lin: [] for lin in lineage_buckets}
    buckets["OTHER"] = []

    for name in names:
        row = global_df[global_df["Product Name*"] == name].iloc[0]
        ptype = str(row["Product Type*"]).lower().strip()
        # determine “lin” exactly the same way you do for colors:
        if ptype in {"flower","pre-roll","concentrate","infused pre-roll","solventless concentrate","vape cartridge"}:
            lin = row["Lineage"].upper()
        elif ptype == "paraphernalia":
            lin = "PARAPHERNALIA"
        elif str(row["Product Strain"]) == "CBD Blend":
            lin = "CBD"
        elif str(row["Product Strain"]) == "Mixed":
            lin = "MIXED"
        else:
            lin = row["Lineage"].upper()

        buckets.setdefault(lin if lin in buckets else "OTHER", []).append(name)

    # now render, bucket-by-bucket in that fixed order
    for lin in lineage_buckets + ["OTHER"]:
        for name in sorted(buckets[lin]):
            bg = LINEAGE_COLOR_MAP.get(lin, "#FFFFFF")
            fg = "white" if bg != "#FFFFFF" else "black"

            var = tkmod.BooleanVar(value=True)
            chk = tkmod.Checkbutton(
                selected_tags_container,
                text=name,
                variable=var,
                bg=bg, fg=fg,
                selectcolor=bg,
                activebackground=bg,
                activeforeground=fg,
                anchor="w", bd=0, highlightthickness=0
            )
            chk.tag_name = name
            chk.pack(fill="x", padx=5, pady=2)
            selected_tags_vars[name] = var
            splash.destroy()

    selected_canvas.update_idletasks()
    selected_canvas.configure(scrollregion=selected_canvas.bbox("all"))
    selected_canvas.yview_moveto(0)

# --- New Section: Selected/Available Tags with "Select All" in Selected Tags ---
selected_tags_all_var = None  # Initialize later in main()

def update_selected_tags_all_state():
    global selected_tags_vars, selected_tags_all_var
    for tag, var in selected_tags_vars.items():
        var.set(selected_tags_all_var.get())

def create_selected_header():
    global selected_tags_all_var
    header_frame = tkmod.Frame(selected_tags_container, bg="lightgray")
    header_frame.pack(fill="x", padx=2, pady=2)
    select_all_chk = tkmod.Checkbutton(header_frame,
                                         text="Select All (Selected Tags)",
                                         variable=selected_tags_all_var,
                                         bg="lightgray",
                                         font=("Arial", 12),
                                         anchor="w",
                                         command=update_selected_tags_all_state)
    select_all_chk.pack(side="left", padx=5)

def move_to_selected():
    splash = show_splash2(root)
    global available_tags_vars, selected_tags_vars, undo_stack

    # 1) Find checked tags in Available
    moved_tags = [tag for tag, var in available_tags_vars.items() if var.get()]

    # 2) Move each to Selected
    for tag in moved_tags:
        var = available_tags_vars.pop(tag)
        selected_tags_vars[tag] = var

    # 3) Re-render both lists (placeholder logic lives in those functions)
    populate_available_tags(list(available_tags_vars.keys()))
    
    populate_selected_tags(list(selected_tags_vars.keys()))

    # 4) Add a divider with current filter values if none exists
    divider_exists = any(getattr(w, "is_divider", False)
                         for w in selected_tags_container.winfo_children())
    if not divider_exists:
        filter_values = []
        if vendor_filter_var.get() != "All":
            filter_values.append("Vendor: " + vendor_filter_var.get())
        if product_brand_filter_var.get() != "All":
            filter_values.append("Brand: " + product_brand_filter_var.get())
        if product_type_filter_var.get() != "All":
            filter_values.append("Type: " + product_type_filter_var.get())
        if lineage_filter_var.get() != "All":
            filter_values.append("Lineage: " + lineage_filter_var.get())
        if product_strain_filter_var.get() != "All":
            filter_values.append("Ratio: " + product_strain_filter_var.get())
        if weight_filter_var.get() != "All":
            filter_values.append("Weight: " + weight_filter_var.get())

        if not filter_values:
            filter_values.append("All")

        divider_text = "------- Selected Filter Values: " + ", ".join(filter_values) + " -------"
        header_divider = tkmod.Label(
            selected_tags_container,
            text=divider_text,
            font=("Arial", 10, "italic"),
            fg="blue",
            bg="lightgray"
        )
        header_divider.is_divider = True
        header_divider.pack(
            fill="x",
            pady=2,
            before=selected_tags_container.winfo_children()[0]
        )

    # 5) Record the move for undo
    if moved_tags:
        undo_stack.append(moved_tags)


    divider_exists = any(getattr(widget, "is_divider", False) for widget in selected_tags_container.winfo_children())
    if not divider_exists:
        filter_values = []
        if vendor_filter_var.get() != "All":
            filter_values.append("Vendor: " + vendor_filter_var.get())
        if product_brand_filter_var.get() != "All":
            filter_values.append("Brand: " + product_brand_filter_var.get())
        if product_type_filter_var.get() != "All":
            filter_values.append("Type: " + product_type_filter_var.get())
        if lineage_filter_var.get() != "All":
            filter_values.append("Lineage: " + lineage_filter_var.get())
        if product_strain_filter_var.get() != "All":
            filter_values.append("Ratio: " + product_strain_filter_var.get())
        if weight_filter_var.get() != "All":
            filter_values.append("Weight: " + weight_filter_var.get())

        if not filter_values:
            filter_values.append("All")

        divider_text = "------- Selected Filter Values: " + ", ".join(filter_values) + " -------"
        header_divider = tkmod.Label(selected_tags_container, text=divider_text,
                                     font=("Arial", 10, "italic"), fg="blue", bg="lightgray")
        header_divider.is_divider = True
        header_divider.pack(fill="x", pady=2, before=selected_tags_container.winfo_children()[0])

    if moved_tags:
        undo_stack.append(moved_tags)
    splash.destroy()



def undo_last_move():
    splash = show_splash2(root)
    global undo_stack, available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    if not undo_stack:
        messagebox.showinfo("Undo", "No moves to undo.")
        return
    last_move = undo_stack.pop()  # Get the last list of moved tags
    for tag in last_move:
        # If the tag is currently in selected tags, remove it from there
        if tag in selected_tags_vars:
            # Remove corresponding widget from selected tags container
            for widget in selected_tags_container.winfo_children():
                if hasattr(widget, "tag_name") and widget.tag_name == tag:
                    widget.destroy()
            var = selected_tags_vars.pop(tag)
            # Set its value to True so that it remains selected when moved back
            var.set(True)
            # Re-add the tag to the available tags container
            new_chk = tkmod.Checkbutton(available_tags_container, text=tag, variable=var, bg="white", anchor="w")
            new_chk.tag_name = tag
            new_chk.pack(fill="x", pady=2)
            available_tags_vars[tag] = var
    splash.destroy()

def clear_selected_list():
    splash = show_splash2(root)
    global selected_tags_container, selected_tags_vars, undo_stack
    if selected_tags_container is None:
        logging.warning("Selected tags container is not initialized.")
        return

    # Iterate over a copy of the child widget list
    for widget in list(selected_tags_container.winfo_children()):
        try:
            widget.destroy()
        except Exception as e:
            logging.error("Error destroying widget in clear_selected_list: %s", e)
    # Clear the dictionaries and undo history
    selected_tags_vars.clear()
    undo_stack.clear()

    # Refresh available product names if necessary.
    try:
        update_all_dropdowns()
    except Exception as e:
        logging.error("Error updating dropdowns after clearing selected: %s", e)

    splash.destroy()

def move_to_available():
    splash = show_splash2(root)

    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container

    # Don’t do anything if there’s literally no real selected tags
    to_move = [
        tag for tag, var in selected_tags_vars.items()
        if var.get()
    ]
    if not to_move:
        return

    for tag in to_move:
        # Find & destroy only the real tag widget
        for widget in list(selected_tags_container.winfo_children()):
            if getattr(widget, "is_placeholder", False):
                # skip the placeholder image
                continue
            if getattr(widget, "tag_name", None) == tag:
                widget.destroy()
                break

        # Move its var back to available
        var = selected_tags_vars.pop(tag)
        chk = tkmod.Checkbutton(
            available_tags_container,
            text=tag,
            variable=var,
            bg="white",
            anchor="w"
        )
        chk.tag_name = tag
        chk.pack(fill="x", pady=2)
        available_tags_vars[tag] = var

    # If nothing remains selected, you might want to show the placeholder:
    if not selected_tags_vars:
        populate_selected_tags([])
    splash.destroy()


def move_tag_to_selected(tag):
    splash = show_splash2(root)
    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    # Find and destroy the widget from the available container
    for widget in available_tags_container.winfo_children():
        if getattr(widget, "tag_name", None) == tag:
            widget.destroy()
            break
    # Pop the BooleanVar from available_tags_vars; if none, create a new one.
    var = available_tags_vars.pop(tag, tkmod.BooleanVar(value=True))
    # Create the checkbutton in the selected container.
    new_chk = tkmod.Checkbutton(selected_tags_container, text=tag, variable=var, bg="lightgray", anchor="w")
    new_chk.tag_name = tag
    new_chk.pack(fill="x", pady=2)
    selected_tags_vars[tag] = var
    splash.destroy()

def move_tag_to_available(tag):
    splash = show_splash2(root)
    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    # Find and destroy the widget from the selected container.
    for widget in selected_tags_container.winfo_children():
        if getattr(widget, "tag_name", None) == tag:
            widget.destroy()
            break
    # Create a new BooleanVar for available.
    new_var = tkmod.BooleanVar(value=False)
    chk = tkmod.Checkbutton(available_tags_container, text=tag, variable=new_var, bg="white", anchor="w")
    chk.tag_name = tag
    chk.pack(fill="x", pady=2)
    available_tags_vars[tag] = new_var
    if tag in selected_tags_vars:
        del selected_tags_vars[tag]

    btn_undo = tkmod.Button(button_container, text="↩️ Undo", font=("Arial", 16), command=undo_last_move)
    btn_plus.grid(row=0, column=0, pady=5)
    btn_minus.grid(row=1, column=0, pady=5)
    clear_selected_btn.grid(row=2, column=0, pady=5)
    btn_undo.grid(row=3, column=0, pady=5)
    splash.destroy()

def edit_template(template_type):
    splash = show_splash2(root)
    """
    Opens the specified template file in the system's default application for editing.
    
    Valid template_type values:
       - 'horizontal'
       - 'vertical'
       - 'mini'
      
    """
    template_type = template_type.lower()
    if template_type == 'horizontal':
        path = resource_path("templates/horizontal.docx")
    elif template_type == 'vertical':
        path = resource_path("templates/vertical.docx")
    elif template_type == 'mini':
        path = resource_path("templates/mini.docx")
    else:
        messagebox.showerror("Error", f"Unknown template type: {template_type}")
        return
    splash.destroy()
    open_file(path)


def populate_product_names(sorted_names=None):
    global available_tags_container, selected_tags_container, available_tags_vars, selected_tags_vars, global_df
    # Preserve names already in selected panel.
    current_selected = set(selected_tags_vars.keys())
    df = global_df.copy()
    if product_type_filter_var.get() != "All":
        df = df[df["Product Type*"] == product_type_filter_var.get()]
    if lineage_filter_var.get() != "All":
        df = df[df["Lineage"] == lineage_filter_var.get()]
    if product_brand_filter_var.get() != "All":
        df = df[df["Product Brand"] == product_brand_filter_var.get()]
    if vendor_filter_var.get() != "All":
        df = df[df["Vendor"] == vendor_filter_var.get()]
    if weight_filter_var.get() != "All":
        df = df[df["CombinedWeight"] == weight_filter_var.get()]
    if product_strain_filter_var.get() != "All":
        df = df[df["Product Strain"] == product_strain_filter_var.get()]
    if sorted_names is None:
        names = sorted(df["Product Name*"].dropna().unique())
    else:
        names = sorted(sorted_names)
    # Remove names already selected:
    names = [name for name in names if name not in current_selected]
    for widget in available_tags_container.winfo_children():
        widget.destroy()
    available_tags_vars.clear()
    for name in names:
        var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(available_tags_container, text=name, variable=var, bg="white", anchor="w")
        chk.tag_name = name
        chk.pack(fill="x", pady=2)
        available_tags_vars[name] = var
   
def sort_products_by(column):
    # Your sorting logic here.
    # For example:
    global global_df
    if global_df is None or column not in global_df.columns:
        return
    filtered_df = global_df.copy()
    # (apply additional filters as needed)
    sorted_df = filtered_df.sort_values(by=column, na_position='last')
    # Refresh the product names (which update available/selected tags)
    populate_product_names(sorted_names=sorted_df["Product Name*"].dropna().unique().tolist())

def set_current_canvas(event, canvas):
    global current_canvas
    current_canvas = canvas

def clear_current_canvas(event):
    global current_canvas
    current_canvas = None

def change_lineage():
    import os, datetime, webbrowser
    from concurrent.futures import ThreadPoolExecutor
    from tkinter import ttk
    global global_df, selected_tags_vars, root, file_entry
    splash = show_splash2(root)
    executor = ThreadPoolExecutor(max_workers=1)

    if global_df is None:
        messagebox.showerror("Error", "No Excel file is loaded.")
        return

    # 1) Capture old lineages
    old_map = {
        name: str(global_df.loc[
            global_df["Product Name*"] == name, "Lineage"
        ].iloc[0]).upper()
        for name in selected_tags_vars
    }

    # 2) Define colors & options (including paraphernalia in pink)
    lineage_map = {
        "SATIVA":        ("(S)",     "#E74C3C"),
        "INDICA":        ("(I)",     "#8E44AD"),
        "HYBRID":        ("(H)",     "#27AE60"),
        "HYBRID/SATIVA": ("(S/H)",   "#E74C3C"),
        "HYBRID/INDICA": ("(I/H)",   "#8E44AD"),
        "CBD":           ("(CBD)",   "#F1C40F"),
        "MIXED":         ("(M)",     "#2C3E50"),
        "PARAPHERNALIA": ("(P)",     "#FF69B4"),
    }
    OPTIONS = list(lineage_map.keys())
    LOG_PATH = os.path.expanduser("~/Downloads/lineage_changes_log.csv")

    # 3) Build popup window
    popup = tkmod.Toplevel(root)
    popup.title("Change Lineage")
    popup.geometry("900x800")
    popup.configure(bg="white")

    # — Scrollable Canvas + InnerFrame —
    list_frame = tkmod.Frame(popup, bg="white")
    list_frame.pack(fill="both", expand=True, padx=10, pady=10)
    canvas = tkmod.Canvas(list_frame, bg="white", highlightthickness=0)
    sb     = tkmod.Scrollbar(list_frame, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=sb.set)
    canvas.pack(side="left", fill="both", expand=True)
    sb.pack(side="right", fill="y")

    inner_frame = tkmod.Frame(canvas, bg="white")
    window = canvas.create_window((0,0), window=inner_frame, anchor="nw")
    inner_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )
    canvas.bind("<Configure>", lambda e: canvas.itemconfig(window, width=e.width))

    # Enable scroll tracking
    canvas.bind("<Enter>", lambda e: set_current_canvas(e, canvas))
    canvas.bind("<Leave>", lambda e: clear_current_canvas(e, canvas))
    canvas.bind("<MouseWheel>", lambda e: on_mousewheel(e, canvas))
    canvas.bind("<Button-4>",    lambda e: on_mousewheel(e, canvas))
    canvas.bind("<Button-5>",    lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<MouseWheel>", lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<Button-4>",    lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<Button-5>",    lambda e: on_mousewheel(e, canvas))

    # Track user selections here
    popup_vars = {}

    # 4) Populate each row: shaded label + Combobox
    for name in sorted(selected_tags_vars):
        old_lin = old_map[name]
        # special case: paraphernalia always shows PARAPHERNALIA
        prod_type = str(global_df.loc[
            global_df["Product Name*"] == name, "Product Type*"
        ].iloc[0]).strip().lower()
        if prod_type == "paraphernalia":
            old_lin = "PARAPHERNALIA"

        abbr, bg = lineage_map.get(old_lin, ("", "#BDC3C7"))

        row = tkmod.Frame(inner_frame, bg=bg)
        row.pack(fill="x", pady=2)

        # shaded product name + old abbr
        lbl = tkmod.Label(
            row,
            text=f"{name}  {abbr}",
            bg=bg,
            fg="white",
            font=("Arial", 16, "bold"),
            anchor="w",
            padx=6, pady=4
        )
        lbl.pack(side="left", fill="x", expand=True)

        # dropdown for new lineage
        var = tkmod.StringVar(value=old_lin)
        popup_vars[name] = var

        combo = ttk.Combobox(
            row, textvariable=var,
            values=OPTIONS, state="readonly", width=12
        )
        combo.pack(side="right", padx=6, pady=4)

    # 5) Save Changes & Cancel buttons
    btn_frame = tkmod.Frame(popup, bg="white")
    btn_frame.pack(fill="x", pady=10)
    tkmod.Button(
        btn_frame, text="Save Changes", font=("Arial",12,"bold"),
        bg="white", fg="green", padx=10, pady=5,
        command=lambda: _apply()
    ).pack(side="right", padx=10)
    tkmod.Button(
        btn_frame, text="Cancel", font=("Arial",12),
        command=popup.destroy
    ).pack(side="right")

    def _apply():
        # 1) Build updated df in‑memory & log diffs
        df2 = global_df.copy()
        ts  = datetime.datetime.now().isoformat()
        os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
        with open(LOG_PATH, "a", encoding="utf-8") as log:
            for name, var in popup_vars.items():
                new_lin = var.get().upper()
                old_lin = old_map[name]
                if new_lin != old_lin:
                    df2.loc[
                        df2["Product Name*"] == name, "Lineage"
                    ] = new_lin
                    if new_lin == "MIXED":
                        df2.loc[
                            df2["Product Name*"] == name, "Product Type*"
                        ] = "Mixed"
                    log.write(f"{ts},{name},{old_lin},{new_lin}\n")

        # 2) Background save & reload
        def save_and_reload():
            nowstr = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            out = os.path.expanduser(f"~/Downloads/{nowstr}_LineageUpdated.xlsx")
            df2.to_excel(out, index=False)
            cleaned = preprocess_excel(out)
            newdf = pd.read_excel(cleaned, engine="openpyxl")
            return out, newdf

        future = executor.submit(save_and_reload)

        def on_done(fut):
            try:
                out_path, newdf = fut.result()
            except Exception as e:
                messagebox.showerror("Error", f"Save/Reload failed: {e}")
                return

            # swap in new DataFrame and refresh UI
            global global_df
            global_df = newdf
            populate_filter_dropdowns()
            populate_product_names()
            file_entry.delete(0, "end")
            file_entry.insert(0, out_path)
            messagebox.showinfo(
                "Done",
                f"Saved to:\n{out_path}\n\n"
                f"Log updated at:\n{LOG_PATH}"
            )
            popup.destroy()

        future.add_done_callback(lambda f: root.after_idle(lambda: on_done(f)))

    # Make the popup modal
    popup.grab_set()
    popup.wait_window()
    splash.destroy()


import tkinter as tkmod
from tkinter import ttk, font, colorchooser

def launch_edit_template():
    top = tkmod.Toplevel(root)
    top.title("Edit Template & Font Settings")
    top.geometry("600x500")
    
    nb = ttk.Notebook(top)
    nb.pack(fill="both", expand=True, padx=10, pady=10)
    
    # We'll store settings here:
    font_settings = {
        tmpl: {
            "family": tkmod.StringVar(value="Arial"),
            "size":      tkmod.IntVar(value=12),
            "bold":      tkmod.BooleanVar(value=True),   # <— auto-checked
            "italic": tkmod.BooleanVar(value=False),
            "underline": tkmod.BooleanVar(value=False),
            "color":  tkmod.StringVar(value="#000000"),
        }
        for tmpl in ("Horizontal","Vertical","Mini")
    }
    
    def make_font_tab(name):
        frm = ttk.Frame(nb)
        nb.add(frm, text=name)
        
        setting = font_settings[name]
        
        # Font Family
        ttk.Label(frm, text="Font Family:").grid(row=0, column=0, sticky="w", pady=5)
        fam_combo = ttk.Combobox(frm, textvariable=setting["family"],
                                 values=sorted(font.families()), width=30)
        fam_combo.grid(row=0, column=1, sticky="w", pady=5)
        
        # Font Size
        ttk.Label(frm, text="Base Font Size (pt):").grid(row=1, column=0, sticky="w")
        size_spin = tkmod.Spinbox(frm, from_=6, to=72, textvariable=setting["size"], width=5)
        size_spin.grid(row=1, column=1, sticky="w")
        
        # Bold / Italic / Underline
        b1 = tkmod.Checkbutton(frm, text="Bold",      variable=setting["bold"])
        b2 = tkmod.Checkbutton(frm, text="Italic",    variable=setting["italic"])
        b3 = tkmod.Checkbutton(frm, text="Underline", variable=setting["underline"])
        b1.grid(row=2, column=0, sticky="w", pady=5)
        b2.grid(row=2, column=1, sticky="w", pady=5)
        b3.grid(row=2, column=2, sticky="w", pady=5)
        
        # Color Picker
        def choose_color():
            col = colorchooser.askcolor(setting["color"].get(), parent=frm)[1]
            if col:
                setting["color"].set(col)
                color_btn.config(bg=col)
        ttk.Label(frm, text="Font Color:").grid(row=3, column=0, sticky="w")
        color_btn = tkmod.Button(frm, text="   ", command=choose_color,
                                 bg=setting["color"].get(), width=3)
        color_btn.grid(row=3, column=1, sticky="w")
        
        # layout tweaks
        for c in range(3):
            frm.columnconfigure(c, weight=1)
        return frm
    
    # create one tab per template
    for tpl in ("Horizontal","Vertical","Mini"):
        make_font_tab(tpl)
    
    # at bottom: OK / Cancel
    btn_frame = ttk.Frame(top)
    btn_frame.pack(fill="x", pady=10)
    def on_ok():
        # here you have all the font_settings[...] values
        # e.g. font_settings["Horizontal"]["family"].get(), etc.
        # save them or apply to your template rendering logic
        top.destroy()
    ttk.Button(btn_frame, text="OK",    command=on_ok).pack(side="right", padx=5)
    ttk.Button(btn_frame, text="Cancel",command=top.destroy).pack(side="right")



def show_instructions_popup():
    # Create a new popup window
    popup = tkmod.Toplevel(root)
    popup.title("POSaBit Instructions")
    popup.geometry("1000x800")
    popup.transient(root)
    popup.grab_set()  # Make the popup modal

    # Create a container to hold the text and image side by side
    container = tkmod.Frame(popup, bg="white", padx=10, pady=10)
    container.pack(fill="both", expand=True)

    # Left side: Instruction Text
    text_label = tkmod.Label(
        container,
        text=posabit_instructions,
        fg="gray",
        bg="white",
        font=("Arial", 16),
        justify="left",
        wraplength=350  # Adjust as needed
    )
    text_label.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

    # Right side: The image (assets/step1.png)
    try:
        image_path = os.path.join("assets", "step1.png")
        step1_image = tkmod.PhotoImage(file=image_path)
    except Exception as e:
        logging.error(f"Error loading image at {image_path}: {e}")
        step1_image = None

    if step1_image:
        image_label = tkmod.Label(container, image=step1_image, bg="white")
        image_label.image = step1_image  # keep a reference to avoid garbage collection
        image_label.grid(row=0, column=1, sticky="nsew")
    else:
        image_label = tkmod.Label(container, text="Image not found", bg="white")
        image_label.grid(row=0, column=1, sticky="nsew")

    # Configure grid weights so that columns share available space equally
    container.columnconfigure(0, weight=1)
    container.columnconfigure(1, weight=1)
    container.rowconfigure(0, weight=1)

    # Add a Close button
    close_btn = tkmod.Button(popup, text="Close", font=("Arial", 12), command=popup.destroy)
    close_btn.pack(pady=10)

    # Wait until the popup is closed before returning (modal behavior)
    popup.wait_window()


def simulate_default_upload():
    default_file = get_default_file()  # Make sure this helper function is defined
    if default_file:
        # Set the file_entry widget to the default file path.
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, default_file)
        label_file.config(text=os.path.basename(default_file))
        logging.debug(f"Default file found: {default_file}")
        try:
            # Process the default file as if it were just uploaded
            cleaned_file = preprocess_excel(default_file)
            logging.debug(f"Preprocessed file: {cleaned_file}")
            global global_df
            global_df = pd.read_excel(cleaned_file, engine="openpyxl")
            logging.debug(f"DataFrame loaded. Columns: {global_df.columns.tolist()}")
            logging.debug(global_df.head())
            populate_filter_dropdowns()
            populate_product_names()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process default file: {e}")

def get_default_file():
    """
    Searches the user's Downloads folder for Excel files that start with "A Greener Today"
    and returns the most recently modified file.
    """
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    files = [f for f in os.listdir(downloads_dir)
             if f.startswith("A Greener Today") and f.lower().endswith(".xlsx")]
    if files:
        files_full_paths = [os.path.join(downloads_dir, f) for f in files]
        # Return the most recent file (by modification time)
        return max(files_full_paths, key=os.path.getmtime)
    return None


def show_splash(root):
    splash = tkmod.Toplevel()
    splash.title("Loading...")
    splash.overrideredirect(True)  # Remove window borders
    splash.configure(bg="white")

    # Load the splash image
    try:
        splash_image_path = resource_path("assets/splash.png")
        splash_image = tkmod.PhotoImage(file=splash_image_path)
        width, height = splash_image.width(), splash_image.height()
    except Exception as e:
        logging.error(f"Error loading splash image: {e}")
        width, height = 400, 200  # Fallback size if loading fails
        splash_image = None

    # Center the splash screen
    screen_width = splash.winfo_screenwidth()
    screen_height = splash.winfo_screenheight()
    x = (screen_width // 2) - (width // 2)
    y = (screen_height // 2) - (height // 2)
    splash.geometry(f"{width}x{height}+{x}+{y}")

    if splash_image:
        label = tkmod.Label(splash, image=splash_image, bg="white")
        label.image = splash_image  # Keep reference to avoid garbage collection
    else:
        label = ttk.Label(splash, text="Loading, please wait...", font=("Arial", 16), background="white")

    label.pack(expand=True)

    splash.lift()
    splash.attributes("-topmost", True)
    splash.update()

    return splash

def show_splash2(root):
    splash = tkmod.Toplevel()
    splash.title("Loading...")
    splash.overrideredirect(True)  # Remove window borders
    splash.configure(bg="white")

    # Load the splash image
    try:
        splash_image_path = resource_path("assets/splash2.gif")
        splash_image = tkmod.PhotoImage(file=splash_image_path)
        width, height = splash_image.width(), splash_image.height()
    except Exception as e:
        logging.error(f"Error loading splash image: {e}")
        width, height = 400, 200  # Fallback size if loading fails
        splash_image = None

    # Center the splash screen
    screen_width = splash.winfo_screenwidth()
    screen_height = splash.winfo_screenheight()
    x = (screen_width // 2) - (width // 2)
    y = (screen_height // 2) - (height // 2)
    splash.geometry(f"{width}x{height}+{x}+{y}")

    if splash_image:
        label = tkmod.Label(splash, image=splash_image, bg="white")
        label.image = splash_image  # Keep reference to avoid garbage collection
    else:
        label = ttk.Label(splash, text="Loading, please wait...", font=("Arial", 16), background="white")

    label.pack(expand=True)

    splash.lift()
    splash.attributes("-topmost", True)
    splash.update()

    return splash

def run_full_process_inventory_slips(selected_df):
        if selected_df.empty:
            messagebox.showerror("Error", "No data selected.")
            return

        records = selected_df.to_dict(orient="records")
        pages = []

        for chunk in chunk_records(records, 4):
            tpl = DocxTemplate(INVENTORY_SLIP_TEMPLATE)
            context = {}

            slot_num = 1
            for rec in chunk:
                product_name = rec.get("Product Name*", "")
                barcode      = rec.get("Barcode*", "")
                qty          = rec.get("Quantity Received*", rec.get("Quantity*", ""))

                if not (product_name or barcode or qty):
                    continue

                try:
                    qty = int(float(qty))
                except (ValueError, TypeError):
                    qty = ""

                context[f"Label{slot_num}"] = {
                    "ProductName":      product_name,
                    "Barcode":          barcode,
                    "AcceptedDate":     rec.get("Accepted Date", ""),
                    "QuantityReceived": qty,
                    "Vendor":           rec.get("Vendor", "")
                }
                slot_num += 1

            # fill the rest of the 4 slots with blanks
            for i in range(slot_num, 5):
                context[f"Label{i}"] = {
                    "ProductName":      "",
                    "Barcode":          "",
                    "AcceptedDate":     "",
                    "QuantityReceived": "",
                    "Vendor":           ""
                }

            tpl.render(context)
            buf = BytesIO()
            tpl.save(buf)
            pages.append(Document(buf))

        if not pages:
            messagebox.showerror("Error", "No documents generated.")
            return

# ------------------ MAIN GUI FUNCTION ------------------
def main():
    global root, vendor_filter_var, product_brand_filter_var, product_type_filter_var
    global lineage_filter_var, product_strain_filter_var, weight_filter_var, quantity_filter_var
    global file_entry, label_file
    global selected_tags_all_var, available_tags_all_var, selected_tags_vars
    global current_canvas, available_tags_container, selected_tags_container
    global placeholder_img
    global print_vendor_back_var

    selected_tags_vars = {}

    root = tkmod.Tk()
    try:
        placeholder_img = tkmod.PhotoImage(
            file=resource_path("assets/placeholder.png"),
            master=root
        )
    except Exception:
        # fallback if resource_path failed
        placeholder_img = tkmod.PhotoImage(
            file="assets/placeholder.png",
            master=root
        )
    root.withdraw()  # Hide main GUI initially until loading is done
    
    context_menu = tkmod.Menu(root, tearoff=0)
    for label, sequence in [
        ("Cut", "<<Cut>>"),
        ("Copy", "<<Copy>>"),
        ("Paste", "<<Paste>>"),
        ("Select All", "<<SelectAll>>"),
    ]:
        context_menu.add_command(
            label=label,
            command=lambda seq=sequence: root.focus_get().event_generate(seq)
        )

    def show_context_menu(event):
        widget = event.widget
        # only on Entry/Text/etc — skip other controls if you like
        if isinstance(widget, (tkmod.Entry, tkmod.Text, ttk.Combobox)):
            context_menu.tk_popup(event.x_root, event.y_root)
        return "break"
    
        # ----- install the right-click menu on all text widgets -----
    for cls in ("Entry", "Text", "TCombobox"):
        # standard right-click
        root.bind_class(cls, "<Button-3>", show_context_menu)
        # on macOS two-finger click (sometimes mapped to Button-2)
        root.bind_class(cls, "<Button-2>", show_context_menu)


    splash = show_splash(root)
        # after you create `root = tkmod.Tk()` in main():

    def normalize_columns(df: pd.DataFrame) -> None:
        """
        For each of these expected columns, add a _norm_<col> lowercase, punctuation-stripped
        helper column — but only if the source column actually exists.
        """
        norm_cols = [
            "Product Type*", "Lineage", "Product Brand", "Vendor",
            "Product Strain", "CombinedWeight",
            "Quantity", "Quantity Received*"
        ]
        for col in norm_cols:
            if col in df.columns:
                norm_col = f"_norm_{col}"
                df[norm_col] = (
                    df[col]
                    .fillna("")
                    .astype(str)
                    .str.lower()
                    .str.replace(r"[^\w\s]", " ", regex=True)
                    .str.strip()
                )


    def load_default_file():
        global global_df
        from pathlib import Path

        downloads_dir = Path.home() / "Downloads"
        candidates = sorted(
            downloads_dir.glob("A Greener Today*.xlsx"),
            key=lambda f: f.stat().st_mtime,
            reverse=True
        )
        if candidates:
            # read the newest one
            global_df = pd.read_excel(str(candidates[0]), engine="openpyxl")
            logging.debug("Default file loaded: %s", candidates[0])
        else:
            global_df = pd.DataFrame()
            logging.debug("No default file found.")
        # **normalize right after load**
        normalize_columns(global_df)

    # Load file asynchronously
    from concurrent.futures import ThreadPoolExecutor
    executor = ThreadPoolExecutor(max_workers=1)
    future = executor.submit(load_default_file)

    def check_load_complete():
        if future.done():
            splash.destroy()
            root.deiconify()
            setup_gui(root)
        else:
            splash.after(100, check_load_complete)

    def setup_gui(root):
        root.title("AGT Price Tag Transformer")

        # DPI‑aware scaling
        dpi_scaling = root.winfo_pixels('1i') / 72
        root.tk.call('tk', 'scaling', dpi_scaling)

        # Center and scale GUI
        sw = root.winfo_screenwidth()
        sh = root.winfo_screenheight()
        width  = int(sw * 0.95)
        height = int(sh * 0.95)
        x = (sw - width) // 2
        y = (sh - height) // 2
        root.geometry(f"{width}x{height}+{x}+{y}")

        bind_global_mousewheel(root)

    # Build the main GUI frames and widgets
    main_frame = tkmod.Frame(root, bg="#228B22")
    main_frame.pack(fill="both", expand=True)

    # ---------------- Left Frame: Upload and Filters ----------------
    left_frame = tkmod.Frame(main_frame, bg="#228B22", width=200)
    left_frame.pack(side="left", fill="y", padx=10, pady=10)
    left_frame.pack_propagate(False)

    def upload_file():
        path = filedialog.askopenfilename(
            filetypes=[("Excel Files", "*.xlsx"), ("CSV Files", "*.csv")]
        )
        if not path:
            return

        # read & preprocess
        cleaned = preprocess_excel(path)
        global global_df
        global_df = pd.read_excel(cleaned, engine="openpyxl")
        logging.debug("Uploaded file loaded. Columns: %s", global_df.columns.tolist())

        # **normalize right after load**
        normalize_columns(global_df)

        # refresh all filters & available‐tags panel
        populate_filter_dropdowns()
        populate_product_names()

    btn_upload = tkmod.Button(left_frame, text="Upload Spreadsheet", command=upload_file,
                               bg="#228B22", font=("Arial", 16), height=2)
    btn_upload.pack(pady=20)

    label_file = tkmod.Label(left_frame, text="No file selected", bg="#228B22", fg="white", font=("Arial", 7))
    label_file.pack(pady=5)

    file_entry = tkmod.Entry(left_frame, bd=0, bg="white", fg="#000716", font=("Arial", 8))
    #file_entry.pack(fill="x", padx=5, pady=5)

    def get_json_url():
            url = json_url_entry.get().strip()
            if not url.lower().startswith("http"):
                messagebox.showerror("Invalid URL", "Please paste a valid JSON URL.")
                return
            # start matching in background…
            threading.Thread(target=_fetch_and_match, args=(url,), daemon=True).start()

            try:
                with urllib.request.urlopen(url) as resp:
                    payload = json.loads(resp.read().decode())
            except Exception as e:
                messagebox.showerror("Error", f"Failed to fetch JSON:\n{e}")
                return

            items       = payload.get("inventory_transfer_items", [])
            vendor_meta = f"{payload.get('from_license_number','')} – {payload.get('from_license_name','')}"
            raw_date    = payload.get("est_arrival_at", "").split("T")[0]

            records = []
            for itm in items:
                records.append({
                    "Product Name*":      itm.get("product_name", ""),
                    "Barcode*":           itm.get("inventory_id", ""),
                    "Quantity Received*": itm.get("qty", ""),
                    "Accepted Date":      raw_date,
                    "Vendor":             vendor_meta,
                })

            df = pd.DataFrame(records)
            run_full_process_inventory_slips(df)

        # --- new JSON section ---
    json_url_entry = tkmod.Entry(left_frame, font=("Arial", 12))
    json_url_entry.pack(fill="x", padx=5, pady=5)

    btn_json = tkmod.Button(
        left_frame,
        text="▶ Load JSON & Match",
        command= get_json_url,
        bg="white", fg="#228B22", font=("Arial", 14)
    )
    btn_json.pack(fill="x", padx=5, pady=(0,10))

        

    # Pre-populate the file_entry if a default file is found.
    default_file = get_default_upload_file()
    if default_file:
        file_entry.insert(0, default_file)
        label_file.config(text=os.path.basename(default_file))

    filter_defs = [
        ("\nVendor:", "vendor_filter_var", "vendor_option"),
        ("\nBrand:", "product_brand_filter_var", "product_brand_option"),
        ("\nProduct Type:", "product_type_filter_var", "product_type_option"),
        ("\nLineage (S/H/I):", "lineage_filter_var", "lineage_option"),
        ("\nCBD Blend:", "product_strain_filter_var", "product_strain_option"),
        ("\nWeight:", "weight_filter_var", "weight_option")
    ]
    for text, var_name, option_name in filter_defs:
        lbl = tkmod.Label(left_frame, text=text, bg="#228B22", font=("Arial", 16), fg="white")
        lbl.pack(pady=3)
        globals()[var_name] = tkmod.StringVar(left_frame, value="All")
        opt = tkmod.OptionMenu(left_frame, globals()[var_name], "All")
        opt.config(bg="white", width=10)
        opt["menu"].config(bg="white")
        opt.pack(pady=5, fill="x")
        globals()[option_name] = opt

    if platform.system() == "Darwin":
        check_font = ("Arial", 10)
        pady_val = 10
    else:
        check_font = ("Segoe UI", 8)
        pady_val = 10

    quantity_filter_var = tkmod.BooleanVar(value=True)
    quantity_chk = tkmod.Checkbutton(left_frame, text="Only show products with Quantity > 0",
                                     variable=quantity_filter_var, bg="#228B22", font=check_font,
                                     fg="white", selectcolor="#228B22", activebackground="#228B22",
                                     activeforeground="white", highlightthickness=0, anchor="w", padx=5)
    quantity_chk.pack(pady=pady_val, fill="x")

    file_entry = tkmod.Entry(left_frame, bd=0, bg="white", fg="#000716")

    def clear_filters():
        splash = show_splash2
        # reset all dropdowns to “All”
        for var in (vendor_filter_var, product_brand_filter_var,
                    product_type_filter_var, lineage_filter_var,
                    product_strain_filter_var, weight_filter_var):
            var.set("All")
        # clear any JSON override
        global json_matched_names
        json_matched_names = []
        json_url_entry.delete(0, "end")
        # *then* rebuild everything from the full sheet
        update_all_dropdowns()
        splash.destroy()
        




    btn_clear = tkmod.Button(left_frame, text="Clear Filter", command=clear_filters,
                              bg="#228B22", font=("Arial", 16), height=4)
    btn_clear.pack(pady=10, fill="x")

    # ---------------- Center Frame: Tag Panels ----------------
    center_frame = tkmod.Frame(main_frame, bg="green", width=420, height=800)
    center_frame.pack(side="left", padx=10, pady=10, fill="x", expand=True)
    center_frame.pack_propagate(False)

    
    # Container for tag panels and move buttons
    tags_frame = tkmod.Frame(center_frame, bg="green")
    tags_frame.pack(fill="both", expand=True)

    # ---- Available Tags Panel (Left) ----
    available_panel = tkmod.Frame(tags_frame, bg="white", width=400)
    available_panel.pack(side="left", fill="both", expand=True)
  
    available_label = tkmod.Label(available_panel, text="Available Tag List:", bg="white", font=("Arial", 14))
    available_label.pack(pady=5)

    sort_buttons_frame = tkmod.Frame(available_panel, bg="#D3D3D3")
    sort_buttons_frame.pack(fill="x", padx=5, pady=5)

    available_header = tkmod.Frame(available_panel, bg="white")
    available_header.pack(fill="x", padx=5, pady=(0,5))
    available_tags_all_var = tkmod.BooleanVar(root, value=True)
    available_select_all_chk = tkmod.Checkbutton(
        available_header,
        text="Select All (Available)",
        variable=available_tags_all_var,
        bg="white",
        font=("Arial", 12),
        anchor="w",
        command=update_available_tags_all_state_available
    )
    available_select_all_chk.pack(side="left", padx=5)

    global available_canvas
    available_canvas = tkmod.Canvas(available_panel, bg="white")
    available_canvas.pack(side="left", fill="both", expand=True)
    available_scrollbar = tkmod.Scrollbar(available_panel, orient="vertical", command=available_canvas.yview)
    available_scrollbar.pack(side="right", fill="y")
    available_canvas.configure(yscrollcommand=available_scrollbar.set)
    available_tags_container = tkmod.Frame(available_canvas, bg="white")
    available_tags_container.bind("<Configure>", lambda event: available_canvas.configure(scrollregion=available_canvas.bbox("all")))
    available_canvas.create_window((0, 0), window=available_tags_container, anchor="nw")
    available_tags_container.bind(
    "<Configure>",
    lambda e: available_canvas.configure(scrollregion=available_canvas.bbox("all"))
)

    available_canvas.bind("<Enter>", lambda event: set_current_canvas(event, available_canvas))
    available_canvas.bind("<Leave>", lambda event: clear_current_canvas(event))
    available_canvas.bind("<MouseWheel>", lambda event: on_mousewheel(event, available_canvas))
    available_canvas.bind("<Button-4>", lambda event: available_canvas.yview_scroll(-1, "units"))
    available_canvas.bind("<Button-5>", lambda event: available_canvas.yview_scroll(1, "units"))

    # ---- Move Buttons Panel (Middle) ----
    move_btn_frame = tkmod.Frame(tags_frame, bg="green", width=100, height=800)
    move_btn_frame.pack(side="left", fill="both", padx=5)
    move_btn_frame.pack_propagate(False)
    button_container = tkmod.Frame(move_btn_frame, bg="green")
    button_container.place(relx=0.5, rely=0.5, anchor="center")
    btn_plus = tkmod.Button(button_container, text=">", font=("Arial", 16), command=move_to_selected)
    btn_minus = tkmod.Button(button_container, text="<", font=("Arial", 16), command=move_to_available)
    clear_selected_btn = tkmod.Button(button_container, text="Clear Selected", font=("Arial", 12), command=clear_selected_list)
    btn_undo = tkmod.Button(button_container, text="Undo", font=("Arial", 12), command=undo_last_move)

    # Instructions '?' button directly under Undo
    btn_instructions = tkmod.Button(button_container, text="?", font=("Arial", 16, "bold"),
                                    fg="#228B22", bg="white", relief="raised",
                                    command=show_instructions_popup)

    # Grid layout
    btn_plus.grid(row=0, column=0, pady=15)
    btn_minus.grid(row=1, column=0, pady=15)
    clear_selected_btn.grid(row=2, column=0, pady=15)
    btn_undo.grid(row=3, column=0, pady=15)
    btn_instructions.grid(row=4, column=0, pady=10)  # '?' button placed here

    


    # ---- Selected Tags Panel (Right) ----
    selected_panel = tkmod.Frame(tags_frame, bg="white", width=425)
    selected_panel.pack(side="left", fill="both", expand=True)
   
    selected_label = tkmod.Label(selected_panel, text="Selected Tag List:", bg="white", font=("Arial", 14))
    selected_label.pack(pady=5)
    selected_header_frame = tkmod.Frame(selected_panel, bg="white")
    selected_header_frame.pack(fill="x", padx=5, pady=5)
    selected_tags_all_var = tkmod.BooleanVar(root, value=True)
    select_all_chk = tkmod.Checkbutton(selected_header_frame,
                                       text="Select All (Selected Tags)",
                                       variable=selected_tags_all_var,
                                       bg="white", font=("Arial", 12),
                                       anchor="w",
                                       command=update_selected_tags_all_state)
    select_all_chk.pack(side="left", padx=5)

    global selected_canvas
    selected_canvas = tkmod.Canvas(selected_panel, bg="white")
    selected_canvas.pack(side="left", fill="both", expand=True)
    selected_scrollbar = tkmod.Scrollbar(selected_panel, orient="vertical", command=selected_canvas.yview)
    selected_scrollbar.pack(side="right", fill="y")
    selected_canvas.configure(yscrollcommand=selected_scrollbar.set)
    selected_tags_container = tkmod.Frame(selected_canvas, bg="white")
    selected_tags_container.bind("<Configure>", lambda event: selected_canvas.configure(scrollregion=selected_canvas.bbox("all")))
    selected_canvas.create_window((0, 0), window=selected_tags_container, anchor="nw")
    selected_tags_container.bind(
    "<Configure>",
    lambda e: selected_canvas.configure(scrollregion=selected_canvas.bbox("all"))
)

    selected_canvas.bind("<Enter>", lambda event: set_current_canvas(event, selected_canvas))
    selected_canvas.bind("<Leave>", lambda event: clear_current_canvas(event))
    selected_canvas.bind("<MouseWheel>", lambda event: on_mousewheel(event, selected_canvas))
    selected_canvas.bind("<Button-4>", lambda event: selected_canvas.yview_scroll(-1, "units"))
    selected_canvas.bind("<Button-5>", lambda event: selected_canvas.yview_scroll(1, "units"))

    # ---------------- Right Frame: Action Buttons ----------------
    right_frame = tkmod.Frame(main_frame, bg="#228B22", width=150)
    right_frame.pack(side="left", fill="y", padx=10, pady=10)
    right_frame.pack_propagate(False)
        # ─── New: Print Vendor to Back checkbox ─────────────────────────
    print_vendor_back_var = tkmod.BooleanVar(value=False)
    vendor_back_chk = tkmod.Checkbutton(
        right_frame,
        text="Print Vendor to Back",
        variable=print_vendor_back_var,
        bg="#228B22",
        fg="white",
        selectcolor="#228B22",
        font=("Arial", 12),
        anchor="w"
    )
    vendor_back_chk.pack(pady=10, fill="x")

       # ─── Scale Factor slider ─────────────────────────
    scale_factor_var = tkmod.DoubleVar(value=1.0)  # default = 1×

    def on_scale_change(val):
        new_scale = scale_factor_var.get()
        # store it somewhere global or pass it into process_chunk
        global SCALE_FACTOR
        SCALE_FACTOR = new_scale
        # (re-run any previews if you like)


    tkmod.Label(
        right_frame,
        text="Font Scale Factor",
        bg="#228B22", fg="white",
        font=("Arial", 12)
    ).pack(pady=(10,0))

        # right after you create your Scale…
    scale_factor_var = tkmod.DoubleVar(value=1.0)
    scale_slider = tkmod.Scale(
        right_frame,
        variable=scale_factor_var,
        from_=0.5, to=2.0,
        resolution=0.05,
        orient="horizontal",
        length=200,
        bg="#228B22",
        fg="white",
        troughcolor="#BBBBBB",
        highlightthickness=0,
        command=on_scale_change
    )
    scale_slider.pack(pady=(0,10))

    # add a Reset button immediately below (or beside) the slider:
    reset_btn = tkmod.Button(
        right_frame,
        text="Reset Scale",
        font=("Arial", 10),
        command=lambda: (
            scale_factor_var.set(1.0),
            on_scale_change(1.0)  # if you want to reapply immediately
        )
    )
    reset_btn.pack(pady=(0,20))






    btn_horizontal = tkmod.Button(
                            right_frame,
                            text="▭ Horizontal Tags",
                            command=lambda: run_full_process_by_group("horizontal"),
                            bg="#228B22", font=("Arial", 16), height=4,
    anchor="w", padx=10
    )
    btn_horizontal.pack(pady=20, fill="x")

    btn_vertical = tkmod.Button(right_frame, text="▯ Vertical Tags",
                                 command=lambda: run_full_process_by_group("vertical"),
                                 bg="#228B22", font=("Arial", 16), height=3)
    btn_vertical.pack(pady=20, fill="x")
    btn_mini = tkmod.Button(right_frame, text="⬜ Mini Tags",
                            command=run_full_process_mini,
                            bg="#228B22", font=("Arial", 16), height=4)
    btn_mini.pack(pady=20, fill="x")

    btn_edit_template = tkmod.Button(right_frame, 
                                    text="🖊️ Edit Template", 
                                    command=launch_edit_template,
                                    bg="#228B22", font=("Arial", 16), height=4)
    btn_edit_template.pack(pady=20, fill="x")
    btn_edit_data = tkmod.Button(right_frame, text="🎨 Fix Lineage",
                                  command=change_lineage,
                                  bg="#228B22", font=("Arial", 16), height=4)
    btn_edit_data.pack(pady=20, fill="x")

    def bind_dropdown_traces():
        vendor_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_brand_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_type_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        lineage_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_strain_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        weight_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
    bind_dropdown_traces()

    from pathlib import Path
    downloads_dir = Path.home() / "Downloads"
    # Use glob to get matching files (case-insensitive if needed)
    matching_files = sorted(downloads_dir.glob("A Greener Today*.xlsx"),
                            key=lambda f: f.stat().st_mtime,
                            reverse=True)
    if matching_files:
        default_path = str(matching_files[0])
        # Update the file_entry widget with the default file path
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, default_path)
        try:
            global_df = pd.read_excel(default_path, engine="openpyxl")
            populate_filter_dropdowns()
            populate_product_names()  # This function should repopulate available tags automatically
            logging.debug("Default file loaded: " + default_path)
        except Exception as e:
            logging.error("Error reading default file: " + str(e))
    else:
        logging.debug("No default file matching 'A Greener Today*.xlsx' found in Downloads.")
        
    simulate_default_upload()
    populate_filter_dropdowns()
    if "Product Name*" not in global_df.columns:
        messagebox.showerror("Missing Column", "'Product Name*' column not found in your uploaded file.")
        return
    populate_product_names()
    # if no tags have been moved yet, show placeholders in both panels:
    populate_selected_tags([])
    check_load_complete()

    logging.debug("Entering mainloop")
    root.mainloop()
    logging.debug("After mainloop (should not reach here until window is closed)")

if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        with open("error.log", "w") as f:
            f.write(traceback.format_exc())
        messagebox.showerror("Application Error", "An error occurred. Please check the error.log file for details.")
        sys.exit(1)